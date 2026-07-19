"""
agents/documentation_agent.py
───────────────────────────────
Phase 8 — Documentation Agent

Builds a comprehensive Model Development Document (.docx) following an 18-section
structure aligned to SR 11-7 / RBI / ECOA expectations, plus three companion
Excel workbooks and an embedded SHAP beeswarm image. Everything is pulled from
`state` — nothing is hardcoded — so it works on Dataset 2 as well.
"""

import os
import json
from datetime import datetime

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.base_agent import BaseAgent
from core.state import PipelineState

try:
    from core.llm import ask, ask_with_usage, CREDIT_RISK_SYSTEM
except Exception:  # pragma: no cover
    ask, ask_with_usage, CREDIT_RISK_SYSTEM = None, None, ""

try:
    from agents.validation_agent import get_rag, KPI_THRESHOLDS
except Exception:  # pragma: no cover
    get_rag, KPI_THRESHOLDS = None, {}


# RAG → cell fill (hex, no #) for Word table shading
RAG_FILL = {"GREEN": "C6EFCE", "AMBER": "FFEB9C", "RED": "FFC7CE", "N/A": "FFFFFF"}

EXCEL_FILES = {
    "dqr":       "DQR_Full_Report.xlsx",
    "features":  "Feature_Selection_Detail.xlsx",
    "models":    "Model_Comparison.xlsx",
    "calib":     "Calibration_Reliability.xlsx",
}


class DocumentationAgent(BaseAgent):

    def __init__(self, output_dir: str = "outputs", verbose: bool = True):
        super().__init__("DocumentationAgent", verbose)
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    # ═══════════════════════════════════════════════════════════════
    def run(self, state: PipelineState) -> PipelineState:
        self._token_usage = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}
        self._log("Rendering SHAP beeswarm image …")
        self._shap_png = self._make_shap_png(state)

        self._log("Generating report charts …")
        state = self._generate_charts(state)

        self._log("Generating companion Excel workbooks …")
        self._excel_paths = self._generate_excels(state)

        self._log("Building Model Development Document (.docx) …")
        state = self._generate_docx(state)

        self._log("Writing plain-text report …")
        state = self._generate_txt(state)

        # Record accumulated LLM token usage before the audit trail is serialised.
        state.llm_token_usage[self.name] = self._token_usage

        self._log("Saving audit trail JSON …")
        state = self._save_audit(state)

        state.agent_responses[self.name] = self.build_response(
            summary="Model Development Document + companion workbooks generated",
            observations=[
                f"Word document: {os.path.basename(state.model_report_path or '')}",
                f"Excel companions: {len(self._excel_paths)} workbooks",
                f"SHAP image embedded: {'yes' if self._shap_png else 'no'}",
                f"Sections: 18 (Exec Summary → Appendices)",
            ],
            reasoning=(
                "All agent outputs aggregated into an 18-section Model Development "
                "Document with GOOD/BAD/INDETERMINATE target definition, 3-way "
                "Dev/Test/OOT metrics, RAG-graded KPIs, calibration, SHAP, and "
                "companion Excel workbooks referenced throughout."
            ),
            recommendations=[
                "Circulate the Word document + Excel appendices for governance sign-off",
                "File the audit-trail JSON alongside the champion model artefact",
            ],
        )
        return state

    # ───────────────────────── helpers ─────────────────────────────
    def _llm(self, prompt: str, max_tokens: int = 350) -> str:
        if ask_with_usage is None:
            return ""
        try:
            text, usage = ask_with_usage(prompt, system=CREDIT_RISK_SYSTEM, max_tokens=max_tokens)
            # Accumulate across the (multiple) narrative calls this agent makes.
            acc = getattr(self, "_token_usage", None)
            if acc is not None:
                acc["input_tokens"] += usage["input_tokens"]
                acc["output_tokens"] += usage["output_tokens"]
                acc["total_tokens"] += usage["total_tokens"]
            return text
        except Exception:
            return ""

    def _load_checkpoints(self) -> dict:
        path = os.path.join(self.output_dir, "checkpoints.json")
        if os.path.exists(path):
            try:
                with open(path, encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    @staticmethod
    def _shade(cell, fill_hex: str):
        """Shade a table cell background."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)

    def _rag_for(self, metric_key: str, value):
        """Deterministic RAG via validation thresholds; ('N/A','') if unknown."""
        if get_rag is None or value in (None, "N/A", ""):
            return "N/A", ""
        try:
            return get_rag(metric_key, float(value))
        except Exception:
            return "N/A", ""

    # ───────────────────────── SHAP image ──────────────────────────
    def _make_shap_png(self, state: PipelineState):
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
        except Exception as e:
            state.log_warning(self.name, f"matplotlib unavailable for SHAP image: {e}")
            return None
        sf = os.path.join(self.output_dir, "models", f"{state.run_id}_shap_sample.json")
        if not os.path.exists(sf):
            return None
        try:
            with open(sf, encoding="utf-8") as f:
                sp = json.load(f)
            names = sp["feature_names"]
            shap = np.array(sp["shap_values"], dtype=float)
            feat = np.array(sp["feature_values"], dtype=float)
            order = np.argsort(np.abs(shap).mean(0))[::-1][:10]
            fig, ax = plt.subplots(figsize=(8, 5))
            rng = np.random.default_rng(42)
            for row, fi in enumerate(reversed(list(order))):
                sv, xv = shap[:, fi], feat[:, fi]
                norm = (xv - np.nanmin(xv)) / (np.nanmax(xv) - np.nanmin(xv) + 1e-8)
                yj = row + rng.uniform(-0.3, 0.3, len(sv))
                sc = ax.scatter(sv, yj, c=norm, cmap="RdBu_r", s=9, alpha=0.6)
            ax.set_yticks(range(len(order)))
            ax.set_yticklabels([names[i] for i in reversed(list(order))])
            ax.axvline(0, color="#333", lw=0.9, ls="--")
            ax.set_xlabel("SHAP value (impact on model output)")
            ax.set_title("SHAP Summary — Feature Impact on Default Prediction")
            cbar = fig.colorbar(sc, ax=ax, ticks=[0, 1])
            cbar.ax.set_yticklabels(["Low", "High"])
            cbar.set_label("Feature value")
            fig.tight_layout()
            path = os.path.join(self.output_dir, f"{state.run_id}_shap_beeswarm.png")
            fig.savefig(path, dpi=120, bbox_inches="tight")
            plt.close(fig)
            self._info(f"SHAP beeswarm image saved → {path}")
            return path
        except Exception as e:
            state.log_warning(self.name, f"SHAP image render failed: {e}")
            return None

    # ───────────────────────── report charts ───────────────────────
    def _generate_charts(self, state: PipelineState) -> PipelineState:
        """Generate 7 report charts (dark theme, headless). Each is isolated in its
        own try/except so one failure never blocks the rest."""
        state.chart_paths = {}
        self._chart_status = {}
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
        except Exception as e:
            state.log_warning(self.name, f"matplotlib unavailable — all charts skipped: {e}")
            self._chart_status = {f"chart{n}": f"skipped: matplotlib unavailable" for n in range(7)}
            return state

        chart_dir = os.path.join(self.output_dir, "charts", state.run_id)
        os.makedirs(chart_dir, exist_ok=True)
        vm = state.validation_metrics or {}

        def _dark(fig, axes):
            fig.patch.set_facecolor("#07090e")
            for a in (axes if isinstance(axes, (list, tuple)) else [axes]):
                a.set_facecolor("#0f1420")
                a.tick_params(colors="#e0e7ff")
                for sp in a.spines.values():
                    sp.set_color("#2a3350")
                a.title.set_color("#e0e7ff")
                a.xaxis.label.set_color("#e0e7ff")
                a.yaxis.label.set_color("#e0e7ff")

        def _save(fig, key):
            path = os.path.join(chart_dir, f"{key}.png")
            fig.savefig(path, dpi=120, bbox_inches="tight", facecolor=fig.get_facecolor())
            plt.close(fig)
            state.chart_paths[key] = path
            self._chart_status[key] = "ok"

        # 1 — Target class balance pie
        try:
            wdf = state.working_df
            good = int((wdf["target"] == 0).sum()) if wdf is not None and "target" in wdf.columns else 0
            bad  = int((wdf["target"] == 1).sum()) if wdf is not None and "target" in wdf.columns else 0
            indet = sum(v.get("count", 0) for v in (state.indeterminate_values or []))
            parts = [(l, v, c) for l, v, c in
                     [("GOOD (0)", good, "#10b981"), ("BAD (1)", bad, "#ef4444"),
                      ("INDETERMINATE", indet, "#6b7280")] if v > 0]
            if not parts:
                raise ValueError("no target counts available")
            fig, ax = plt.subplots(figsize=(6, 5))
            ax.pie([p[1] for p in parts], labels=[p[0] for p in parts], colors=[p[2] for p in parts],
                   autopct=lambda pct: f"{pct:.1f}%", textprops={"color": "#e0e7ff"})
            ax.set_title("Chart 4 — Target Class Balance")
            _dark(fig, ax)
            _save(fig, "chart4_target")
        except Exception as e:
            self._chart_status["chart4_target"] = f"skip: {e}"
            state.log_warning(self.name, f"chart4_target failed: {e}")

        # 2 — Univariate AUC top 10 (behavioral leakage)
        try:
            blf = state.behavioral_leakage_flags or {}
            if not blf:
                raise ValueError("no behavioral leakage flags")
            items = sorted(blf.items(), key=lambda x: x[1])[-10:]
            fig, ax = plt.subplots(figsize=(7, 5))
            ax.barh([k for k, _ in items], [v for _, v in items], color="#4f6ef7")
            ax.axvline(0.65, color="#f59e0b", ls="--", lw=1)
            ax.set_xlabel("Univariate AUC")
            ax.set_title("Chart 7 — Univariate AUC (Top 10)")
            _dark(fig, ax)
            _save(fig, "chart7_univariate_auc")
        except Exception as e:
            self._chart_status["chart7_univariate_auc"] = f"skip: {e}"
            state.log_warning(self.name, f"chart7 failed: {e}")

        # 3 — IV for all candidate features (color by strength)
        try:
            iv = state.iv_table
            if not isinstance(iv, pd.DataFrame) or iv.empty:
                raise ValueError("no IV table")
            d = iv.sort_values("iv")
            scol = {"Strong": "#10b981", "Medium": "#4f6ef7", "Weak": "#6b7280",
                    "Suspicious": "#f59e0b", "Useless": "#ef4444"}
            colors = [scol.get(str(s).split()[0], "#4f6ef7") for s in d.get("strength", ["Medium"] * len(d))]
            fig, ax = plt.subplots(figsize=(7, max(4, len(d) * 0.24)))
            ax.barh(d["feature"].astype(str), d["iv"], color=colors)
            ax.set_xlabel("Information Value")
            ax.set_title("Chart 8 — IV by Feature (colour = strength)")
            ax.tick_params(axis="y", labelsize=6)
            _dark(fig, ax)
            _save(fig, "chart8_iv_all")
        except Exception as e:
            self._chart_status["chart8_iv_all"] = f"skip: {e}"
            state.log_warning(self.name, f"chart8 failed: {e}")

        # 4 — Correlation heatmap of final selected features
        try:
            sel = [f for f in (state.selected_features or [])
                   if state.engineered_df is not None and f in state.engineered_df.columns]
            if len(sel) < 2:
                raise ValueError("need ≥2 selected features")
            corr = state.engineered_df[sel].corr()
            fig, ax = plt.subplots(figsize=(1.2 + 0.55 * len(sel), 1.2 + 0.55 * len(sel)))
            im = ax.imshow(corr.values, cmap="RdBu_r", vmin=-1, vmax=1)
            ax.set_xticks(range(len(sel))); ax.set_xticklabels(sel, rotation=90, fontsize=7)
            ax.set_yticks(range(len(sel))); ax.set_yticklabels(sel, fontsize=7)
            cb = fig.colorbar(im, ax=ax, fraction=0.046)
            cb.ax.tick_params(colors="#e0e7ff")
            ax.set_title("Chart 9 — Correlation (selected features)")
            _dark(fig, ax)
            _save(fig, "chart9_correlation")
        except Exception as e:
            self._chart_status["chart9_correlation"] = f"skip: {e}"
            state.log_warning(self.name, f"chart9 failed: {e}")

        # 5 — ROC curve for champion
        try:
            from sklearn.metrics import roc_curve, roc_auc_score
            model, X_te, y_te = state.champion_model, state.X_test, state.y_test
            if model is None or X_te is None or y_te is None:
                raise ValueError("champion model / test set unavailable")
            yp = model.predict_proba(X_te)[:, 1]
            fpr, tpr, _ = roc_curve(y_te, yp)
            auc = roc_auc_score(y_te, yp)
            fig, ax = plt.subplots(figsize=(6, 5))
            ax.plot(fpr, tpr, color="#4f6ef7", lw=2, label=f"AUC = {auc:.4f}")
            ax.plot([0, 1], [0, 1], color="#6b7280", ls="--", lw=1)
            ax.set_xlabel("False Positive Rate"); ax.set_ylabel("True Positive Rate")
            ax.set_title("Chart 11 — ROC Curve")
            leg = ax.legend(facecolor="#0f1420", edgecolor="#2a3350", labelcolor="#e0e7ff")
            _dark(fig, ax)
            _save(fig, "chart11_roc")
        except Exception as e:
            self._chart_status["chart11_roc"] = f"skip: {e}"
            state.log_warning(self.name, f"chart11 failed: {e}")

        # 6 — Calibration curve before / after
        try:
            cal = state.calibration_results or {}
            ru, rc = cal.get("reliability_uncalibrated", []), cal.get("reliability_calibrated", [])
            if not ru or not rc:
                raise ValueError("no calibration reliability tables")
            fig, ax = plt.subplots(figsize=(6, 5))
            ax.plot([0, 1], [0, 1], color="#6b7280", ls="--", lw=1, label="Perfect")
            ax.plot([r["mean_predicted"] for r in ru], [r["actual_rate"] for r in ru],
                    marker="o", color="#ef4444", label="Uncalibrated")
            ax.plot([r["mean_predicted"] for r in rc], [r["actual_rate"] for r in rc],
                    marker="o", color="#10b981", label=f"Calibrated ({cal.get('recommended_method', '')})")
            ax.set_xlabel("Mean predicted probability"); ax.set_ylabel("Observed default rate")
            ax.set_title("Chart 14 — Calibration Curve")
            ax.legend(facecolor="#0f1420", edgecolor="#2a3350", labelcolor="#e0e7ff")
            _dark(fig, ax)
            _save(fig, "chart14_calibration")
        except Exception as e:
            self._chart_status["chart14_calibration"] = f"skip: {e}"
            state.log_warning(self.name, f"chart14 failed: {e}")

        # 7 — Default rate by decile
        try:
            dec = vm.get("decile_table", [])
            if not dec:
                raise ValueError("no decile table")
            fig, ax = plt.subplots(figsize=(7, 5))
            ax.bar([str(d.get("decile")) for d in dec], [d.get("bad_rate", 0) for d in dec],
                   color="#4f6ef7")
            ax.set_xlabel("Score Decile"); ax.set_ylabel("Default Rate")
            ax.set_title("Chart 17 — Default Rate by Decile")
            _dark(fig, ax)
            _save(fig, "chart17_decile")
        except Exception as e:
            self._chart_status["chart17_decile"] = f"skip: {e}"
            state.log_warning(self.name, f"chart17 failed: {e}")

        # 23 — Fairness: predicted default rate by proxy group (first attribute with results)
        try:
            fr = state.fairness_results or {}
            attr = next((a for a, g in fr.items() if g), None)
            if attr is None:
                raise ValueError("no fairness results")
            groups = fr[attr]
            names = list(groups.keys())
            preds = [groups[n]["mean_predicted"] for n in names]
            overall = groups[names[0]]["mean_predicted"] - groups[names[0]]["diff_from_avg"]
            cmap = {"High": "#ef4444", "Medium": "#f59e0b", "Low": "#10b981"}
            colors = [cmap.get(groups[n]["concern_level"], "#4f6ef7") for n in names]
            fig, ax = plt.subplots(figsize=(max(6, len(names) * 0.55), 5))
            ax.bar([str(n) for n in names], preds, color=colors)
            ax.axhline(overall, color="#9ca3af", ls="--", lw=1, label=f"Overall {overall:.1%}")
            ax.set_ylabel("Mean predicted default prob")
            ax.set_title(f"Chart 23 — Fairness: predicted default by {attr}")
            ax.tick_params(axis="x", rotation=45)
            ax.legend(facecolor="#0f1420", edgecolor="#2a3350", labelcolor="#e0e7ff")
            _dark(fig, ax)
            _save(fig, "chart23_fairness")
        except Exception as e:
            self._chart_status["chart23_fairness"] = f"skip: {e}"
            state.log_warning(self.name, f"chart23 failed: {e}")

        _ok = [k for k, v in self._chart_status.items() if v == "ok"]
        self._info(f"Report charts: {len(_ok)}/{len(self._chart_status)} generated → {chart_dir}")
        for k, v in self._chart_status.items():
            if v != "ok":
                self._info(f"  {k}: {v}")
        return state

    # ───────────────────────── Excel companions ────────────────────
    def _generate_excels(self, state: PipelineState) -> dict:
        rid = state.run_id
        paths = {}

        # 1) DQR_Full_Report — variable profile, consistency checks, PSI/CSI
        try:
            p = os.path.join(self.output_dir, f"{rid}_{EXCEL_FILES['dqr']}")
            with pd.ExcelWriter(p) as xl:
                prof = [{"Column": c, **{k: v.get(k) for k in
                        ("dtype", "semantic_type", "role", "n_unique", "missing_pct", "business_meaning")}}
                        for c, v in (state.schema_profile or {}).items()]
                (pd.DataFrame(prof) if prof else pd.DataFrame({"note": ["no schema"]})).to_excel(
                    xl, sheet_name="Variable_Profile", index=False)
                cc = state.dqr_report.get("consistency_checks", [])
                (pd.DataFrame(cc) if cc else pd.DataFrame({"note": ["no rules applicable"]})).to_excel(
                    xl, sheet_name="Consistency_Checks", index=False)
                csi = state.validation_metrics.get("csi_results", {})
                csi_rows = [{"Feature": k, "CSI": (v.get("csi") if isinstance(v, dict) else v),
                             "Assessment": (v.get("assessment") if isinstance(v, dict) else "")}
                            for k, v in csi.items()]
                psi = state.psi_results or {}
                psi_rows = [{"Metric": "PSI (score)", "Value": psi.get("psi_score"),
                             "Assessment": psi.get("assessment"), "Split": psi.get("split_label")}]
                pd.DataFrame(psi_rows + csi_rows).to_excel(xl, sheet_name="Stability_PSI_CSI", index=False)
            paths["dqr"] = p
        except Exception as e:
            state.log_warning(self.name, f"DQR excel failed: {e}")

        # 2) Feature_Selection_Detail — IV table, correlation pairs, waterfall
        try:
            p = os.path.join(self.output_dir, f"{rid}_{EXCEL_FILES['features']}")
            with pd.ExcelWriter(p) as xl:
                iv = state.iv_table
                (iv if isinstance(iv, pd.DataFrame) else pd.DataFrame(iv or [])).to_excel(
                    xl, sheet_name="IV_Table", index=False)
                pairs = state.dqr_report.get("high_correlation_pairs", [])
                (pd.DataFrame(pairs) if pairs else pd.DataFrame({"note": ["no pairs > threshold"]})).to_excel(
                    xl, sheet_name="Correlation_Pairs", index=False)
                rej = [{"Feature": k, "Reason": v} for k, v in (state.rejected_features or {}).items()]
                (pd.DataFrame(rej) if rej else pd.DataFrame({"note": ["none"]})).to_excel(
                    xl, sheet_name="Rejected_Features", index=False)
                pd.DataFrame({"Selected_Feature": state.selected_features}).to_excel(
                    xl, sheet_name="Selected_Features", index=False)
            paths["features"] = p
        except Exception as e:
            state.log_warning(self.name, f"Feature excel failed: {e}")

        # 3) Model_Comparison — 5 models, hyperparameters, calibration reliability
        try:
            p = os.path.join(self.output_dir, f"{rid}_{EXCEL_FILES['models']}")
            with pd.ExcelWriter(p) as xl:
                rows = []
                for name, m in (state.model_metrics or {}).items():
                    rows.append({"Model": name, "Champion": name == state.champion_model_name,
                                 "AUC_train": m.get("auc_train"), "AUC_test": m.get("auc_test"),
                                 "KS": m.get("ks"), "Gini": m.get("gini"), "F1": m.get("f1"),
                                 "Overfit": m.get("overfit")})
                (pd.DataFrame(rows) if rows else pd.DataFrame({"note": ["no models"]})).to_excel(
                    xl, sheet_name="Model_Comparison", index=False)
                champ_m = (state.model_metrics or {}).get(state.champion_model_name, {})
                hp = champ_m.get("best_params", {})
                (pd.DataFrame([{"Hyperparameter": k, "Value": v} for k, v in hp.items()])
                 if hp else pd.DataFrame({"note": ["defaults used"]})).to_excel(
                    xl, sheet_name="Champion_Hyperparams", index=False)
                cal = state.calibration_results or {}
                rc = cal.get("reliability_calibrated", [])
                (pd.DataFrame(rc) if rc else pd.DataFrame({"note": ["no calibration"]})).to_excel(
                    xl, sheet_name="Calibration_Reliability", index=False)
                # Confusion matrix — 3-way (Train/Test/OOT) + Test classification report
                _cms = {"Train": state.validation_metrics.get("confusion_matrix_train"),
                        "Test":  state.validation_metrics.get("confusion_matrix_test")
                                 or state.validation_metrics.get("confusion_matrix"),
                        "OOT":   state.validation_metrics.get("confusion_matrix_oot")}
                _cms = {k: v for k, v in _cms.items() if v}
                if _cms:
                    _keys = [("N (sample size)", "n"), ("True Positive", "true_positive"),
                             ("False Positive", "false_positive"), ("True Negative", "true_negative"),
                             ("False Negative", "false_negative"), ("Accuracy", "accuracy"),
                             ("Precision (BAD)", "precision"), ("Recall (BAD)", "recall"),
                             ("Specificity (GOOD)", "specificity"), ("F1 (BAD)", "f1_score")]
                    _cm_df = pd.DataFrame(
                        [{"Metric": lbl, **{s: cm.get(key) for s, cm in _cms.items()}} for lbl, key in _keys])
                    _cm_df.to_excel(xl, sheet_name="Confusion_Matrix", index=False, startrow=0)
                    _cr = (_cms.get("Test") or {}).get("classification_report", {})
                    _cr_rows = [{"Class (Test)": c, "Precision": _cr[c].get("precision"),
                                 "Recall": _cr[c].get("recall"), "F1": _cr[c].get("f1-score"),
                                 "Support": _cr[c].get("support")}
                                for c in ("Good (0)", "Bad (1)") if isinstance(_cr.get(c), dict)]
                    if _cr_rows:
                        pd.DataFrame(_cr_rows).to_excel(
                            xl, sheet_name="Confusion_Matrix", index=False, startrow=len(_cm_df) + 3)
                else:
                    pd.DataFrame({"note": ["no confusion matrix"]}).to_excel(
                        xl, sheet_name="Confusion_Matrix", index=False)
            paths["models"] = p
        except Exception as e:
            state.log_warning(self.name, f"Model excel failed: {e}")

        self._info(f"Companion Excel workbooks: {len(paths)}")
        return paths

    # ═══════════════════════════ DOCX ══════════════════════════════
    def _generate_docx(self, state: PipelineState) -> PipelineState:
        doc = Document()
        rid = state.run_id
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        vm  = state.validation_metrics or {}
        cps = self._load_checkpoints()

        base = doc.styles["Normal"]
        base.font.name, base.font.size = "Calibri", Pt(10)

        def h1(t):
            p = doc.add_heading(t, level=1)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        def h2(t): doc.add_heading(t, level=2)
        def body(t): doc.add_paragraph(str(t))
        def bullet(t): doc.add_paragraph(str(t), style="List Bullet")

        def table(headers, rows, rag_col=None):
            """rows: list[list]; rag_col: index whose value name (GREEN/AMBER/RED) shades the cell."""
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            for i, hd in enumerate(headers):
                c = t.rows[0].cells[i]
                c.text = str(hd)
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.bold = True
            for r in rows:
                cells = t.add_row().cells
                for i, val in enumerate(r):
                    cells[i].text = "" if val is None else str(val)
                if rag_col is not None and rag_col < len(r):
                    tag = str(r[rag_col]).upper().split()[0] if r[rag_col] else "N/A"
                    self._shade(cells[rag_col], RAG_FILL.get(tag, "FFFFFF"))
            doc.add_paragraph()
            return t

        def ref(txt):
            p = doc.add_paragraph()
            run = p.add_run(f"Reference: {txt}")
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        def chart(key, caption):
            """Embed a generated chart PNG (if present) with an italic caption."""
            p = (state.chart_paths or {}).get(key)
            if p and os.path.exists(p):
                try:
                    doc.add_picture(p, width=Inches(6))
                except Exception:
                    return
                cp = doc.add_paragraph()
                r = cp.add_run(caption)
                r.italic = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # ── Cover ──
        title = doc.add_heading("Model Development Document", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f"Credit Risk Model — {state.dataset_name}\nRun ID: {rid}\nGenerated: {now}")
        doc.add_page_break()

        # ── overall verdict from KPI scoreboard ──
        scoreboard = [r for r in vm.get("kpi_scoreboard", [])
                      if r.get("KPI") != "Gini Gap (Dev→OOT)"]
        rags = [str(r.get("RAG", "")).upper() for r in scoreboard]
        n_red, n_amber = rags.count("RED"), rags.count("AMBER")
        verdict = "GREEN" if (n_red == 0 and n_amber == 0) else ("AMBER" if n_red == 0 else "RED")

        # ═══ 1. Executive Summary ═══
        h1("1. Executive Summary")
        ctx = self._llm(
            f"Write a 120-word business-context paragraph for a credit-risk model "
            f"development report. Dataset: {state.dataset_name}. Champion: "
            f"{state.champion_model_name}. Purpose: predict borrower default to support "
            f"underwriting. Overall verdict: {verdict}. Be factual, governance-appropriate."
        )
        body(ctx or (f"This document records the development of a credit-risk model on the "
                     f"{state.dataset_name} dataset. The champion model, {state.champion_model_name}, "
                     f"was selected to rank borrowers by probability of default in support of "
                     f"underwriting decisions."))

        h2("1.1 Overall Model Verdict")
        vp = doc.add_paragraph()
        vr = vp.add_run(f"{verdict}")
        vr.bold = True
        vr.font.size = Pt(14)
        vr.font.color.rgb = {"GREEN": RGBColor(0x10, 0x80, 0x40),
                             "AMBER": RGBColor(0xB0, 0x70, 0x00),
                             "RED":   RGBColor(0xC0, 0x20, 0x20)}[verdict]
        body(f"Derived from the KPI scoreboard: {rags.count('GREEN')} GREEN, "
             f"{n_amber} AMBER, {n_red} RED across {len(rags)} discrimination/stability KPIs.")

        h2("1.2 Key Results Summary")
        champ_m = (state.model_metrics or {}).get(state.champion_model_name, {})
        cal = state.calibration_results or {}
        key_rows = []

        def kpi_row(name, key, value):
            rag, lbl = self._rag_for(key, value)
            key_rows.append([name, "N/A" if value in (None, "") else value, rag, lbl])

        kpi_row("AUC / AUROC (Test)", "auc", vm.get("auc"))
        kpi_row("Gini (Test)", "gini", vm.get("gini"))
        kpi_row("KS (Test)", "ks", vm.get("ks"))
        if vm.get("auc_oot") is not None:
            kpi_row("AUC (OOT)", "auc", vm.get("auc_oot"))
        kpi_row("PSI (score stability)", "psi", (state.psi_results or {}).get("psi_score"))
        # calibration ECE (no threshold in get_rag → LLM/manual)
        ece = (cal.get("uncalibrated") or {}).get("ece")
        ece_rag = "GREEN" if isinstance(ece, (int, float)) and ece <= 0.02 else ("RED" if ece is not None else "N/A")
        key_rows.append(["ECE (uncalibrated)", ece if ece is not None else "N/A", ece_rag,
                         "≤0.02 target"])
        table(["KPI", "Value", "RAG", "Note"], key_rows, rag_col=2)
        ref(f"KPI scoreboard & thresholds — {EXCEL_FILES['models']} and Appendix 18.1")

        h2("1.3 Validation KPIs — meaning")
        body("‘Validation KPIs’ here means the model-performance metrics measured on held-out "
             "data (Test and OOT): discrimination (AUC/Gini/KS), stability (PSI/CSI), and "
             "calibration (Brier/ECE). Each carries a RAG badge against the thresholds in "
             "Appendix 18.1. RAG grades absent from source data are computed from those thresholds.")

        h2("1.4 Summary of Human Review Decisions")
        # Only the FORMAL checkpoint keys (cp1/cp2/cp3) are reported here. Page-level
        # overrides (data_understanding, feature_overrides, champion_override, …) are
        # NOT formal checkpoint decisions and are deliberately ignored.
        try:
            run_dt = datetime.strptime(rid, "RUN_%Y%m%d_%H%M%S")
        except Exception:
            run_dt = datetime.now()

        def _stale_note(ts):
            """Flag checkpoints recorded > 24h before this run (likely a prior session)."""
            try:
                if (run_dt - datetime.fromisoformat(ts)).total_seconds() > 24 * 3600:
                    return " (from earlier session — verify still applicable)"
            except Exception:
                pass
            return ""

        hr_rows = []
        for cpk, label in [("cp1", "Checkpoint 1 — Target Definition"),
                           ("cp2", "Checkpoint 2 — Feature Shortlist"),
                           ("cp3", "Checkpoint 3 — Model Sign-Off")]:
            c = cps.get(cpk, {})
            ts = c.get("timestamp", "")
            hr_rows.append([label, c.get("decision", "Pending"),
                            (c.get("analyst_name") or c.get("analyst_notes") or "—"),
                            f"{ts or '—'}{_stale_note(ts)}"])
        table(["Checkpoint", "Decision", "Analyst / Notes", "Timestamp"], hr_rows)

        # ═══ 2. Project Team ═══
        h1("2. Project Team")
        table(["Role", "Name", "Responsibility"],
              [["Model Developer", "[to be completed]", "Development & documentation"],
               ["Model Validator", "[to be completed]", "Independent validation"],
               ["Model Owner", "[to be completed]", "Business ownership & sign-off"],
               ["Governance/MRM", "[to be completed]", "Risk oversight"]])

        # ═══ 3. Purpose and Scope ═══
        h1("3. Purpose and Scope")
        body("Purpose: develop a statistical model that ranks borrowers by probability of "
             "default (PD) to support origination/underwriting decisions.")
        body("Standards basis: RBI model-governance expectations, US Federal Reserve/OCC SR 11-7 "
             "(Model Risk Management), and ECOA/Reg-B (fair-lending, adverse-action requirements).")
        body(f"Scope: single champion model ({state.champion_model_name}) trained and evaluated on "
             f"{state.dataset_name}; out-of-time and blind-dataset scoring supported.")

        # ═══ 4. Data Overview ═══
        h1("4. Data Overview")
        wdf = state.working_df
        if wdf is not None:
            body(f"Observations (after indeterminate exclusion): {wdf.shape[0]:,} | "
                 f"Raw columns: {wdf.shape[1]}")
            if "target" in wdf.columns:
                body(f"Overall default rate: {wdf['target'].mean():.2%}")
        cats = {"Numeric": len(state.numeric_columns), "Categorical": len(state.categorical_columns),
                "Date": len(state.date_columns), "Identifier": len(state.id_columns),
                "Structural leakage": len(state.leakage_columns)}
        table(["Variable Category", "Count"], [[k, v] for k, v in cats.items()])
        ref(f"Full variable profile — {EXCEL_FILES['dqr']} → Variable_Profile")

        # ═══ 5. Data Quality Assessment ═══
        h1("5. Data Quality Assessment")
        dup = state.dqr_report.get("duplicates", {})
        body(f"Duplicate loan IDs: {dup.get('duplicate_ids', 0)} | "
             f"Duplicate member IDs: {dup.get('duplicate_members', 0)}")

        h2("5.1 Missing Value Analysis (top 10)")
        miss = state.missing_summary or {}
        top_miss = sorted([(k, v) for k, v in miss.items()
                           if (v.get("n_missing", 0) if isinstance(v, dict) else 0) > 0],
                          key=lambda x: -(x[1].get("pct_missing", 0) if isinstance(x[1], dict) else 0))[:10]
        table(["Column", "Missing Count", "Missing %"],
              [[c, v.get("n_missing", ""), f"{v.get('pct_missing',0)*100:.1f}%"]
               for c, v in top_miss] or [["—", "—", "—"]])

        h2("5.2 Consistency Check Results")
        body("Each rule encodes a business-logic constraint. The ‘Violations’ column is the "
             "COUNT OF ROWS that fail that rule (e.g. funded_amnt > loan_amnt). Zero violations "
             "= PASS. Rules whose columns are absent in this dataset are simply not applied.")
        cc = state.dqr_report.get("consistency_checks", [])
        table(["Rule", "Description", "Violations", "Status"],
              [[r.get("Rule"), r.get("Description"), r.get("Violations"),
                r.get("Status")] for r in cc] or [["—", "—", "—", "—"]], rag_col=3)
        _ccs = state.dqr_report.get("consistency_summary", {})
        body(_ccs.get("note", ""))
        ref(f"{EXCEL_FILES['dqr']} → Consistency_Checks")

        # (Variable Stability PSI/CSI moved to §12.2 — it requires a trained model + OOT.)

        # ═══ 6. Development, Test & OOT Data Plan ═══
        h1("6. Development, Test & OOT Data Plan")
        sd = state.split_details or {}
        h2("6.1 Data Split Strategy")
        body(f"Method: {sd.get('method', state.split_method)}")
        split_rows = []
        tot = sd.get("total")
        for label, size_k, dr_k in [("Total", "total", None), ("Dev", "dev_size", "dev_default_rate"),
                                    ("Train", "train_size", "train_default_rate"),
                                    ("Test", "test_size", "test_default_rate"),
                                    ("OOT", "oot_size", "oot_default_rate")]:
            sz = sd.get(size_k)
            if sz is None:
                continue
            pct = f"{sz/tot*100:.0f}%" if tot else "—"
            dr = sd.get(dr_k)
            split_rows.append([label, f"{sz:,}", pct, f"{dr*100:.2f}%" if isinstance(dr, (int, float)) else "—"])
        table(["Set", "Size", "% of Total", "Default Rate"], split_rows or [["—", "—", "—", "—"]])
        body("Random seed: 42 (fixed for reproducibility) | Split: OOT = latest 20% by issue_year "
             "(deterministic) | Train/Test = 70/30 stratified random split within remaining 80% (seed=42)")
        if sd.get("input_file_hash"):
            body(f"Input file SHA-256: {sd['input_file_hash'][:16]}… (recorded in the split manifest; an "
                 "unchanged input file reproduces an identical split).")
        if vm.get("auc_oot") is None and not sd.get("oot_size"):
            body("Note: OOT is required for full validation. Where no in-sample OOT is carved, "
                 "OOT is evaluated on blind Dataset 2 via evaluate_oot.py.")

        h2("6.2 Target Class Distribution")
        self._target_distribution_table(state, table)

        h2("6.3 Observation Window")
        body("Development uses the earliest vintages; the out-of-time window uses the most recent "
             "vintages by origination date (issue_year) to test temporal robustness.")

        # ═══ 7. Data Preparation Waterfall ═══
        h1("7. Data Preparation Waterfall")
        h2("7.1 Row Waterfall — Observation Exclusions")
        raw_rows = state.raw_df.shape[0] if state.raw_df is not None else None
        indet = sum(v.get("count", 0) for v in (state.indeterminate_values or []))
        final_rows = wdf.shape[0] if wdf is not None else None
        table(["Stage", "Rows", "Change", "Reason"],
              [["1. Raw dataset", f"{raw_rows:,}" if raw_rows else "—", "", "As supplied"],
               ["2. Exclude INDETERMINATE", f"{(raw_rows-indet):,}" if raw_rows else "—",
                f"-{indet:,}", "Censored / unknown-outcome loans excluded"],
               ["3. Modelling population", f"{final_rows:,}" if final_rows else "—", "", "GOOD + BAD only"]])

        # (Column/Feature Waterfall moved to §10.1 — it belongs with Variable Selection.)

        h2("7.2 Indeterminate Exclusions Detail")
        if state.indeterminate_values:
            table(["Excluded Value", "Count", "Reason"],
                  [[v.get("value"), f"{v.get('count', 0):,}", v.get("reason", "indeterminate outcome")]
                   for v in state.indeterminate_values])
        else:
            body("No indeterminate values were excluded for this dataset.")

        # ═══ 8. Target Variable Definition ═══
        h1("8. Target Variable Definition")
        body(state.target_definition or "")
        body("The target is defined in THREE categories — GOOD, BAD and INDETERMINATE — not a "
             "simple binary. INDETERMINATE loans (censored or unknown future status) are EXCLUDED "
             "from training because their true outcome is not yet observable.")
        self._target_mapping_table(state, table)
        chart("chart4_target", "Chart 4 — Target Class Balance. GOOD / BAD / INDETERMINATE split "
                               "of the population (indeterminates excluded from modelling).")

        # ═══ 9. Feature Engineering ═══
        h1("9. Feature Engineering")
        body(state.dqr_report.get("feature_engineering_summary") or
             "Derived features created from raw fields; raw columns superseded by engineered ones dropped.")
        h2("9.1 Derived Features")
        fl = state.feature_log or []
        table(["Feature", "Rationale"], [[e.get("feature", ""), e.get("rationale", "")]
              for e in fl[:25]] or [["—", "—"]])

        h2("9.2 Categorical Encoding — One-Hot Encoding")
        body("Categorical variables with N unique values are converted into N binary flag columns, "
             "each indicating presence of that category. This avoids imposing false ordinal "
             "relationships that would exist if categories were simply numbered (e.g. encoding "
             "purpose as 1..K would wrongly imply an order/magnitude between purposes).")

        h2("9.3 Behavioral Leakage Detection")
        body("Each numeric field is scored by its univariate AUC against the target; fields with "
             "unusually high standalone AUC are flagged as potential leakage and reviewed.")
        blf = state.behavioral_leakage_flags or {}
        table(["Column", "Univariate AUC", "Flag"],
              [[c, f"{a:.3f}" if isinstance(a, (int, float)) else a,
                "⚠ review" if isinstance(a, (int, float)) and a >= 0.65 else "ok"]
               for c, a in sorted(blf.items(), key=lambda x: -(x[1] if isinstance(x[1], (int, float)) else 0))[:12]]
              or [["—", "—", "—"]])
        chart("chart7_univariate_auc", "Chart 7 — Univariate AUC (Top 10). Columns whose standalone "
                                       "AUC ≥ 0.65 are flagged for behavioral-leakage review.")

        # ═══ 10. Variable Selection ═══
        h1("10. Variable Selection")
        body(state.dqr_report.get("variable_selection_rationale") or
             "Features selected by Information Value with a correlation filter; suspicious IV (>0.50) flagged.")
        _n_raw = len(state.schema_profile or {})
        _n_sel = len(state.selected_features or [])
        h2(f"10.1 Feature Reduction Journey ({_n_raw} → {_n_sel})")
        self._feature_journey_table(state, table)
        ref(f"{EXCEL_FILES['features']} → IV_Table / Rejected_Features")
        chart("chart8_iv_all", "Chart 8 — Information Value for all candidate features, coloured by "
                               "strength (green = strong, red = useless, amber = suspicious).")
        h2("10.2 Final Selected Features (IV with RAG)")
        iv_df = state.iv_table
        rows = []
        if isinstance(iv_df, pd.DataFrame) and not iv_df.empty:
            sel = iv_df[iv_df["feature"].isin(state.selected_features)].sort_values("iv", ascending=False)
            for _, r in sel.iterrows():
                rag, _ = self._rag_for("iv", r.get("iv"))
                rows.append([r.get("feature"), f"{r.get('iv', 0):.4f}", r.get("strength", ""), rag])
        table(["Feature", "IV", "Strength", "RAG"], rows or [["—", "—", "—", "—"]], rag_col=3)
        body(f"Selected: {len(state.selected_features)} | Rejected: {len(state.rejected_features)}")
        ref(f"{EXCEL_FILES['features']} → IV_Table, Correlation_Pairs")
        chart("chart9_correlation", "Chart 9 — Correlation heatmap of the final selected features "
                                    "(red = positive, blue = negative correlation).")

        # ═══ 11. Model Development ═══
        h1("11. Model Development")
        h2("11.1 Algorithm Comparison (all candidates)")
        mm = state.model_metrics or {}
        table(["Model", "AUC (test)", "KS", "Gini", "F1", "Overfit Δ", "Champion"],
              [[n, m.get("auc_test"), m.get("ks"), m.get("gini"), m.get("f1"),
                m.get("overfit"), "★" if n == state.champion_model_name else ""]
               for n, m in mm.items()] or [["—"]*7])
        body(f"Champion rationale: {state.model_selection_rationale or 'selected on best overfit-penalised test AUC.'}")
        h2("11.2a Hyperparameter Search Process")
        _hp_narr, _hp_search = self._hp_search_narrative(state)
        body(_hp_narr)
        if _hp_search:
            table(["Parameter", "Search Range"],
                  [[r["Parameter"], r["Search Range"]] for r in _hp_search])
        _oth = state.optuna_trials_history or {}
        if state.champion_model_name == "XGBoost" and _oth.get("top_3_trials"):
            _best = _oth.get("best_auc") or 0
            _worst = _oth.get("worst_auc") or 0
            body(f"Best trial: AUC={_best} | Worst trial: AUC={_worst} | "
                 f"Improvement: {round(_best - _worst, 4)} across {_oth.get('n_trials')} trials.")
            table(["Trial #", "CV AUC", "Params"],
                  [[t.get("trial"), t.get("auc"),
                    ", ".join(f"{k}={v}" for k, v in (t.get("params") or {}).items())]
                   for t in _oth["top_3_trials"]])
        h2("11.2b Selected Hyperparameters — Final Values")
        hp = champ_m.get("best_params", {})
        table(["Hyperparameter", "Value"], [[k, v] for k, v in hp.items()] or [["defaults", "—"]])
        h2("11.3 Reproducibility")
        bullet("Random seed fixed at 42 across split, model training and Optuna (TPESampler seed=42).")
        bullet("Class imbalance handled via scale_pos_weight / class_weight='balanced'.")
        bullet("Library versions pinned in requirements.txt; audit trail JSON captures the run.")
        ref(f"{EXCEL_FILES['models']} → Model_Comparison, Champion_Hyperparams")

        # ═══ 12. Model Results & Validation ═══
        h1("12. Model Results & Validation")
        h2("12.1 Discrimination — 3-way (Dev/Test/OOT)")
        auc_dev = champ_m.get("auc_train")
        table(["Metric", "Development (train)", "Test", "OOT"],
              [["AUC", auc_dev, vm.get("auc"), vm.get("auc_oot", vm.get("auc_oot_d1"))],
               ["Gini", (2*auc_dev-1) if isinstance(auc_dev, (int, float)) else "—",
                vm.get("gini"), vm.get("gini_oot", vm.get("gini_oot_d1"))],
               ["KS", champ_m.get("ks"), vm.get("ks"), vm.get("ks_oot", vm.get("ks_oot_d1"))]])
        chart("chart11_roc", f"Chart 11 — ROC Curve. AUC = {vm.get('auc', 'N/A')} on the test set.")
        h2("12.2 Variable & Score Stability (PSI / CSI)")
        body("Stability metrics require a trained model and OOT evaluation, hence reported here "
             "rather than during initial data quality review.")
        psi = state.psi_results or {}
        body("PSI (Population Stability Index) measures how much the SCORE distribution shifts "
             "between the development and out-of-time samples. CSI (Characteristic Stability Index) "
             "is the SAME formula applied per input FEATURE, isolating which variables drive any "
             "instability. VDI (Variable Deterioration Index) is NOT yet computed and is recorded "
             "as an open item (see §17).")
        psi_rag, _ = self._rag_for("psi", psi.get("psi_score"))
        table(["Stability Metric", "Value", "Assessment", "RAG"],
              [["PSI (score)", psi.get("psi_score", "N/A"), psi.get("assessment", "—"), psi_rag]]
              + [[f"CSI: {k}", (v.get('csi') if isinstance(v, dict) else v),
                  (v.get('assessment') if isinstance(v, dict) else ""),
                  self._rag_for("csi", (v.get('csi') if isinstance(v, dict) else v))[0]]
                 for k, v in list((vm.get("csi_results") or {}).items())], rag_col=3)
        ref(f"{EXCEL_FILES['dqr']} → Stability_PSI_CSI")
        h2("12.3 Overfit Assessment")
        body(f"Train–Test AUC gap (overfit Δ): {champ_m.get('overfit', 'N/A')} "
             f"(target < 0.03 for GREEN).")
        h2("12.4 KPI Scoreboard (RAG)")
        table(["KPI", "Value", "RAG", "Strength"],
              [[r.get("KPI"), r.get("Value"), r.get("RAG"), r.get("Strength")]
               for r in scoreboard] or [["—"]*4], rag_col=2)
        h2("12.5 Decile Rank-Order Check")
        dec = vm.get("decile_table", [])
        if dec:
            table(["Decile", "N", "Bad Rate", "Avg Prob", "RO Break"],
                  [[d.get("decile"), d.get("n"), f"{d.get('bad_rate',0):.3f}",
                    f"{d.get('avg_prob',0):.3f}", "⚠" if d.get("ro_break") else ""] for d in dec])
        chart("chart17_decile", "Chart 17 — Default rate by score decile (should rise monotonically "
                                "with predicted risk; breaks indicate rank-order issues).")
        h2("12.6 Champion vs Challengers")
        ch = vm.get("challenger_table", [])
        table(["Model", "AUC", "KS", "Gini", "Champion"],
              [[c.get("model"), c.get("auc_test"), c.get("ks"), c.get("gini"), c.get("champion")]
               for c in ch] or [["see 11.1"]*5])
        self._build_confusion_matrix_section(state, doc)

        # ═══ 13. Calibration ═══
        h1("13. Calibration")
        body("Theory: a model can rank risk well (high AUC) yet output probabilities that do not "
             "match observed default rates. Calibration maps raw scores to well-calibrated "
             "probabilities. Platt scaling (sigmoid) fits a logistic transform; isotonic fitting "
             "learns a monotonic step function. Both are MONOTONIC, so ranking/AUC is unchanged — "
             "only the probability values move. Quality is measured by Brier score and Expected "
             "Calibration Error (ECE); we target ECE ≤ 0.02.")
        if cal:
            table(["Variant", "Brier", "ECE", "AUC"],
                  [["Uncalibrated", cal.get("uncalibrated", {}).get("brier"),
                    cal.get("uncalibrated", {}).get("ece"), cal.get("uncalibrated", {}).get("auc")],
                   ["Sigmoid (Platt)", cal.get("sigmoid", {}).get("brier"),
                    cal.get("sigmoid", {}).get("ece"), cal.get("sigmoid", {}).get("auc")],
                   ["Isotonic", cal.get("isotonic", {}).get("brier"),
                    cal.get("isotonic", {}).get("ece"), "unchanged"]])
            chart("chart14_calibration", f"Chart 14 — Calibration curve: uncalibrated vs "
                                         f"{cal.get('recommended_method', 'recommended')} method.")
            body(f"Uncalibrated ECE status: {cal.get('ece_status', '—')}. "
                 f"Recommended method: {cal.get('recommended_method', '—')}.")
            rc = cal.get("reliability_calibrated", [])
            if rc:
                h2("13.1 Post-Calibration Reliability")
                table(["Bin", "Count", "Mean Predicted", "Actual Rate", "Gap"],
                      [[r.get("bin"), r.get("count"), r.get("mean_predicted"),
                        r.get("actual_rate"), r.get("gap")] for r in rc])
            ref(f"{EXCEL_FILES['models']} → Calibration_Reliability")
        else:
            body("Calibration results not available for this run.")

        # ═══ 14. Model Explainability (SHAP) ═══
        h1("14. Model Explainability (SHAP)")
        body(state.shap_summary or "SHAP TreeExplainer applied to the champion model on a sample "
             "of the test set to attribute each prediction to its features.")
        h2("14.1 Global Feature Importance — SHAP Beeswarm")
        if self._shap_png and os.path.exists(self._shap_png):
            doc.add_picture(self._shap_png, width=Inches(6.2))
            body("Red = high feature value, blue = low. Points right of centre increase default "
                 "risk; left decrease it.")
            _ihd = doc.add_paragraph()
            _ihd.add_run("Interpretation").bold = True
            for _line in self._build_shap_analysis(state).split("\n"):
                if _line.strip():
                    body(_line)
        else:
            body("SHAP beeswarm image unavailable (re-run pipeline to regenerate).")
        h2("14.2 Top Features by Mean |SHAP|")
        fi = state.feature_importance or {}
        table(["Feature", "Mean |SHAP|"],
              [[f, f"{v:.5f}"] for f, v in sorted(fi.items(), key=lambda x: -x[1])[:10]] or [["—", "—"]])
        h2("14.3 Individual Prediction Example (Adverse-Action style)")
        adv = state.adverse_action_codes or {}
        first = next(iter(adv.values()), None) if adv else None
        if first:
            table(["Rank", "Feature", "Value", "Contribution"],
                  [[i+1, r.get("feature"), r.get("value"), r.get("contribution", r.get("shap"))]
                   for i, r in enumerate(first.get("top_reasons", [])[:5])] or [["—"]*4])
        else:
            body("Adverse-action example not available.")

        # ═══ 15. Fairness & Bias ═══
        h1("15. Fairness & Bias Assessment")
        body("Methodology: proxy protected attributes (e.g. home ownership, verification status) "
             "are compared for predicted-risk parity on the test set. Note: this dataset lacks direct "
             "protected attributes (race, gender, age); proxy analysis is indicative only and cannot "
             "substitute for a full fair-lending review.")
        fr = state.fairness_results or {}
        if fr:
            h2("15.1 Predicted-Risk Parity by Proxy Attribute")
            _fair_flagged = 0
            fair_rows = []
            for attr, groups in fr.items():
                for g, v in groups.items():
                    if v.get("concern_level") in ("Medium", "High"):
                        _fair_flagged += 1
                    fair_rows.append([attr, g, v.get("count"), v.get("mean_predicted"),
                                      v.get("actual_rate"), v.get("diff_from_avg"),
                                      v.get("concern_level")])
            body(f"{len(fr)} proxy attributes assessed; {_fair_flagged} group(s) flagged "
                 f"(predicted risk deviating > 10 percentage points from the overall mean).")
            table(["Attribute", "Group", "N", "Mean Predicted", "Actual Rate", "Diff from Avg", "Concern"],
                  fair_rows)
            chart("chart23_fairness", "Chart 23 — Predicted default rate by proxy group. Bars beyond "
                                      "the dashed overall-mean line by >10pp indicate potential disparate impact.")
        else:
            body("No fairness results were computed for this run.")
        body("Important caveat: the model must undergo ECOA/Reg-B disparate-impact testing on "
             "protected classes before customer-facing deployment.")

        # ═══ 16. Governance & Sign-Off ═══
        h1("16. Model Governance & Sign-Off")
        h2("16.1 Reviewer Decision")
        cp3 = cps.get("cp3", {})
        table(["Field", "Value"],
              [["Decision", cp3.get("decision", "[pending]")],
               ["Analyst", cp3.get("analyst_name", "[pending]")],
               ["Notes", cp3.get("sign_off_notes", "—")],
               ["Timestamp", cp3.get("timestamp", "—")]])
        h2("16.2 HITL Checkpoints")
        table(["Checkpoint", "Decision"],
              [["1 — Target Definition", cps.get("cp1", {}).get("decision", "Pending")],
               ["2 — Feature Shortlist", cps.get("cp2", {}).get("decision", "Pending")],
               ["3 — Model Sign-Off", cps.get("cp3", {}).get("decision", "Pending")]])
        body("AI/ML governance: every phase was executed by an autonomous agent with human review "
             "gates; all decisions and metrics are recorded in the audit-trail JSON. Regulatory "
             "alignment: SR 11-7 (documentation, validation, monitoring), ECOA (adverse action).")

        # ═══ 17. Assumptions, Limitations, Risks ═══
        h1("17. Assumptions, Limitations, Risks & Caveats")
        alr = self._llm(
            "In <=140 words, list the key assumptions, limitations and risks for a Lending-Club "
            "credit default model that excludes post-origination fields, excludes indeterminate "
            "loans, and is trained on funded loans only. Bullet style."
        )
        if alr:
            body(alr)
        table(["Type", "Item"],
              [["Assumption", "Indeterminate (censored) loans excluded; GOOD/BAD only"],
               ["Assumption", "Imputation applied to missing numeric features"],
               ["Limitation", "Trained on funded loans only — survivorship/selection bias"],
               ["Limitation", "No direct protected attributes — fairness is proxy-based"],
               ["Risk", "Self-reported income/DTI subject to misrepresentation"],
               ["Risk", "Population drift — quarterly PSI/CSI monitoring required"],
               ["Open item", "VDI (Variable Deterioration Index) not yet computed"]])

        # ═══ 18. Appendices ═══
        h1("18. Appendices")
        h2("18.1 KPI Threshold Tables")
        if KPI_THRESHOLDS:
            thr_rows = []
            for m, t in KPI_THRESHOLDS.items():
                thr_rows.append([m, str(t.get("strong_green", "")), str(t.get("acceptable_green", "")),
                                 str(t.get("amber", "")), str(t.get("red", t.get("red_high", "")))])
            table(["Metric", "Strong GREEN", "Acceptable GREEN", "AMBER", "RED"], thr_rows)
        h2("18.2 Attached Excel Workbooks")
        table(["Workbook", "Contents"],
              [[f"{rid}_{EXCEL_FILES['dqr']}", "Variable profile, consistency checks, PSI/CSI"],
               [f"{rid}_{EXCEL_FILES['features']}", "IV table, correlation pairs, selection detail"],
               [f"{rid}_{EXCEL_FILES['models']}", "5-model comparison, hyperparameters, calibration"]])
        h2("18.3 Glossary")
        for term, dfn in [("AUC/AUROC", "Area under ROC curve — rank-ordering power"),
                          ("Gini", "2×AUC−1"),
                          ("KS", "Max separation between cumulative good/bad distributions"),
                          ("IV", "Information Value — predictive strength of a feature"),
                          ("PSI/CSI", "Population/Characteristic Stability Index"),
                          ("ECE", "Expected Calibration Error"),
                          ("SHAP", "Shapley additive explanations — per-prediction attributions"),
                          ("OOT", "Out-of-time validation sample")]:
            bullet(f"{term}: {dfn}")

        # ── Save ──
        docx_path = os.path.join(self.output_dir, f"{rid}_Model_Development_Document.docx")
        doc.save(docx_path)
        state.model_report_path = docx_path
        self._info(f"Model Development Document saved → {docx_path}")
        return state

    # ── section helpers ─────────────────────────────────────────────
    def _target_distribution_table(self, state, table):
        tm = state.target_mapping or {}
        wdf = state.working_df
        good = bad = 0
        if wdf is not None and "target" in wdf.columns:
            bad = int((wdf["target"] == 1).sum())
            good = int((wdf["target"] == 0).sum())
        indet = sum(v.get("count", 0) for v in (state.indeterminate_values or []))
        tot = good + bad + indet
        def pct(n): return f"{n/tot*100:.2f}%" if tot else "—"
        table(["Category", "Count", "% of total"],
              [["GOOD (0)", f"{good:,}", pct(good)],
               ["BAD (1)", f"{bad:,}", pct(bad)],
               ["INDETERMINATE (excluded)", f"{indet:,}", pct(indet)]])

    def _target_mapping_table(self, state, table):
        tm = state.target_mapping or {}
        rows = []
        for raw, code in tm.items():
            cat = "BAD (1)" if code == 1 else "GOOD (0)" if code == 0 else "INDETERMINATE (excluded)"
            rows.append([str(raw), cat])
        for v in (state.indeterminate_values or []):
            if not any(str(v.get("value")) == r[0] for r in rows):
                rows.append([str(v.get("value")), "INDETERMINATE (excluded)"])
        table(["Raw Value", "Mapped Category"], rows or [["—", "—"]])

    def _feature_journey_table(self, state, table):
        schema = state.schema_profile or {}
        total = len(schema)
        leak = len(state.leakage_columns or [])
        ids = len([c for c, v in schema.items() if v.get("role") == "identifier"])
        dates = len([c for c, v in schema.items() if v.get("role") == "date"])
        other = len([c for c, v in schema.items() if v.get("role") == "other"])
        after = total - leak - ids - dates - other
        eng = len(state.feature_log or [])
        rej = state.rejected_features or {}
        iv_rej = len([v for v in rej.values() if any(t in v for t in ("IV", "Useless", "Weak", "Suspicious"))])
        corr_rej = len([v for v in rej.values() if "correlation" in v.lower()])
        final = len(state.selected_features or [])
        table(["Stage", "Columns", "Change"],
              [["1. Raw columns", total, ""],
               ["2. Remove structural leakage", total-leak, f"-{leak}"],
               ["3. Remove identifiers", total-leak-ids, f"-{ids}"],
               ["4. Remove date columns", total-leak-ids-dates, f"-{dates}"],
               ["5. Feature engineering", after+eng, f"+{eng}"],
               ["6. IV filter", after+eng-iv_rej, f"-{iv_rej}"],
               ["7. Correlation filter", after+eng-iv_rej-corr_rej, f"-{corr_rej}"],
               ["8. Final selected", final, ""]])

    def _hp_search_narrative(self, state):
        """Explain HOW the champion's hyperparameters were chosen — dynamically,
        based on the actual champion and the real search space / trial budget."""
        name = state.champion_model_name
        n_trials = state.optuna_trials_run or (state.optuna_trials_history or {}).get("n_trials", "N")
        wdf = state.working_df
        dr = wdf["target"].mean() if (wdf is not None and "target" in wdf.columns) else None
        dr_txt = f"{dr:.1%}" if dr is not None else "the observed"

        if name == "XGBoost":
            # Ranges mirror the actual Optuna objective in model_development_agent._train_xgb.
            search_space = {
                "n_estimators":     "[100, 500]",
                "max_depth":        "[3, 7]",
                "learning_rate":    "[0.01, 0.3] (log scale)",
                "subsample":        "[0.6, 1.0]",
                "colsample_bytree": "[0.5, 1.0]",
                "min_child_weight": "[5, 50]",
            }
            narrative = (
                f"Hyperparameters were tuned using Optuna's Tree-structured Parzen Estimator (TPE) "
                f"sampler over {n_trials} trials, each scored by 3-fold cross-validated AUC on the "
                f"training set. TPE builds a probabilistic model of which parameter regions historically "
                f"produced better AUC and progressively concentrates the search there rather than "
                f"sampling uniformly at random — typically converging faster than grid or random search "
                f"within a limited trial budget. The search space explored was:")
            return narrative, [{"Parameter": k, "Search Range": v} for k, v in search_space.items()]

        if name == "LightGBM":
            return (
                f"LightGBM was trained with fixed, literature-standard defaults rather than an exhaustive "
                f"search, given its role as a fast comparison baseline against the Optuna-tuned XGBoost "
                f"candidate. Parameters mirror common production defaults (100 estimators, 31 leaves, "
                f"0.05 learning rate) with scale_pos_weight computed from the observed class ratio to "
                f"handle the {dr_txt} default rate without oversampling."), []

        if name == "GradientBoosting_AutoML":
            return (
                "GradientBoosting used a fixed, deliberately constrained configuration (100 estimators, "
                "max_depth 3, learning_rate 0.05) as a lightweight ensemble baseline. Adaptive search was "
                "reserved for the XGBoost candidate; the constrained grid here would not justify an "
                "adaptive sampler's typical trial budget."), []

        return (f"{name} was trained with default scikit-learn parameters (no hyperparameter tuning "
                f"applied)."), []

    def _build_shap_analysis(self, state):
        """Dynamic written interpretation of the SHAP beeswarm — concentration of
        importance + per-driver domain narrative, built from actual SHAP values."""
        feat_imp = state.feature_importance or {}
        if not feat_imp:
            return "SHAP analysis not available."

        sorted_feats = sorted(feat_imp.items(), key=lambda x: x[1], reverse=True)
        top5 = sorted_feats[:5]
        total_importance = sum(v for _, v in sorted_feats)
        top_pct = sum(v for _, v in top5) / total_importance * 100 if total_importance else 0

        lines = []
        lines.append(
            "The SHAP beeswarm plot ranks features by their average impact on the model's predicted "
            "default probability. Each point represents one borrower in the test sample; its horizontal "
            "position shows how much that feature pushed the prediction up (increased risk) or down "
            "(decreased risk) for that specific borrower, and its colour shows whether that borrower's "
            "value for the feature was high (red) or low (blue).")
        lines.append(
            f"\nThe top {len(top5)} features together account for approximately {top_pct:.0f}% of total "
            f"feature importance in the model, indicating that risk is driven by a concentrated but not "
            f"single-variable set of signals.")

        interpretations = {
            'int_rate_clean': "reflects the interest rate Lending Club assigned at origination based on the borrower's underwriting risk grade — higher rates were already priced by the platform as higher risk, so this feature's dominance confirms the model is picking up genuine, pre-existing risk signal rather than an artifact.",
            'loan_to_income': "captures how large the requested loan is relative to the borrower's annual income — higher ratios indicate the loan represents a larger financial commitment relative to repayment capacity.",
            'dti': "measures existing debt burden relative to income — borrowers already carrying substantial debt obligations have less capacity to absorb a new loan payment.",
            'annual_inc': "reflects absolute income level — lower income borrowers have a smaller buffer against income shocks.",
            'revol_util': "measures how much of available revolving credit is currently drawn — high utilisation is a classic early-warning indicator of credit stress.",
            'tot_cur_bal': "reflects total existing balances across accounts — a proxy for overall indebtedness.",
            'open_acc_ratio': "reflects the proportion of a borrower's credit accounts that remain open and active — very high or very low ratios can signal unusual credit behaviour.",
            'verification_ordinal': "reflects whether income was independently verified — unverified income carries more uncertainty about the borrower's stated ability to repay.",
            'total_rev_hi_lim': "reflects the total revolving credit limit extended to the borrower by other lenders — a proxy for perceived creditworthiness by the broader market.",
            'loan_amnt': "reflects the absolute size of the requested loan — larger loans carry larger absolute exposure at default.",
        }

        lines.append("\nKey drivers, in order of importance:")
        for i, (feat, val) in enumerate(top5, 1):
            interp = interpretations.get(
                feat, "is one of the model's most influential predictors based on the training data.")
            lines.append(f"{i}. {feat} (mean |SHAP| = {val:.4f}) — {interp}")

        lines.append(
            "\nTaken together, these drivers align with standard credit-risk intuition: pricing signal "
            "(interest rate), affordability (loan-to-income, DTI, income), and credit behaviour "
            "(utilisation, balances) all contribute independently rather than the model relying on any "
            "single dominant signal — supporting robustness against changes in any one input distribution.")

        return "\n".join(lines)

    def _build_confusion_matrix_section(self, state, document):
        """§12.7 — combined Train/Test/OOT confusion matrix + Test classification report."""
        cm_train = state.validation_metrics.get('confusion_matrix_train')
        cm_test  = state.validation_metrics.get('confusion_matrix_test') or \
                   state.validation_metrics.get('confusion_matrix')
        cm_oot   = state.validation_metrics.get('confusion_matrix_oot')
        if not cm_test:
            return
        thr = cm_test.get('threshold', 0.5)

        document.add_heading('12.7 Confusion Matrix & Classification Report', level=2)
        document.add_paragraph(
            "Computed at a 0.5 probability threshold for reporting purposes only — this is NOT "
            "necessarily the correct business decision threshold. The optimal cutoff should be set "
            "based on the actual cost ratio between false negatives and false positives once agreed "
            "with the business. Shown across Train, Test and OOT to check for consistency of "
            "classification performance alongside the ranking metrics reported in Section 12.1.")

        samples = [(l, c) for l, c in
                   [('Train', cm_train), ('Test', cm_test), ('OOT', cm_oot)] if c is not None]

        t = document.add_table(rows=1, cols=len(samples) + 1)
        t.style = 'Table Grid'
        t.rows[0].cells[0].text = 'Metric'
        for i, (label, _) in enumerate(samples, 1):
            t.rows[0].cells[i].text = label
        for metric_label, key, fmt in [
            ('N (sample size)', 'n', '{:,}'),
            ('True Positive', 'true_positive', '{:,}'),
            ('False Positive', 'false_positive', '{:,}'),
            ('True Negative', 'true_negative', '{:,}'),
            ('False Negative', 'false_negative', '{:,}'),
            ('Accuracy', 'accuracy', '{:.4f}'),
            ('Precision (BAD class)', 'precision', '{:.4f}'),
            ('Recall (BAD class)', 'recall', '{:.4f}'),
            ('Specificity (GOOD class)', 'specificity', '{:.4f}'),
            ('F1 Score (BAD class)', 'f1_score', '{:.4f}'),
        ]:
            r = t.add_row().cells
            r[0].text = metric_label
            for i, (_, cm) in enumerate(samples, 1):
                r[i].text = fmt.format(cm.get(key, 0))
        document.add_paragraph()

        gap_note = ""
        if cm_oot:
            recall_gap = abs(cm_test['recall'] - cm_oot['recall'])
            precision_gap = abs(cm_test['precision'] - cm_oot['precision'])
            gap_note = (f" Recall gap Test→OOT: {recall_gap:.4f}; Precision gap: {precision_gap:.4f} — "
                        f"{'consistent' if max(recall_gap, precision_gap) < 0.05 else 'some drift observed'} "
                        f"classification performance across time periods.")
        document.add_paragraph(
            f"Interpretation: at the {thr} threshold, the model correctly identifies "
            f"{cm_test['recall'] * 100:.1f}% of actual defaulters on the Test set (recall), and when it "
            f"predicts default it is correct {cm_test['precision'] * 100:.1f}% of the time (precision)."
            f"{gap_note} Precision and recall both shift as the threshold moves and should be tuned "
            f"jointly with business cost assumptions rather than left at the default 0.5 cutoff.")

        report = cm_test.get('classification_report', {})
        if report:
            document.add_paragraph()
            document.add_paragraph('Full Classification Report — Test set (per class):').runs[0].bold = True
            ct = document.add_table(rows=1, cols=5)
            ct.style = 'Table Grid'
            for i, h in enumerate(['Class', 'Precision', 'Recall', 'F1-Score', 'Support']):
                ct.rows[0].cells[i].text = h
            for cls in ['Good (0)', 'Bad (1)']:
                if cls in report:
                    s = report[cls]
                    r = ct.add_row().cells
                    r[0].text = cls
                    r[1].text = f"{s.get('precision', 0):.4f}"
                    r[2].text = f"{s.get('recall', 0):.4f}"
                    r[3].text = f"{s.get('f1-score', 0):.4f}"
                    r[4].text = f"{int(s.get('support', 0)):,}"

        cap = document.add_paragraph()
        cr_run = cap.add_run("Reference: Model_Comparison.xlsx → Confusion_Matrix")
        cr_run.italic = True
        cr_run.font.size = Pt(8)
        cr_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ───────────────────────── plain text ──────────────────────────
    def _generate_txt(self, state: PipelineState) -> PipelineState:
        rid = state.run_id
        vm = state.validation_metrics or {}
        lines = [
            "=" * 70,
            "MODEL DEVELOPMENT DOCUMENT (text summary)",
            "=" * 70,
            f"Run ID   : {rid}",
            f"Dataset  : {state.dataset_name}",
            f"Champion : {state.champion_model_name}",
            f"AUC={vm.get('auc','N/A')}  KS={vm.get('ks','N/A')}  Gini={vm.get('gini','N/A')}  "
            f"AUC_OOT={vm.get('auc_oot','N/A')}",
            f"Validation: {'PASS' if state.validation_passed else 'CONDITIONAL'}",
            "",
            "Full detail in the Word document and the three Excel companions.",
        ]
        path = os.path.join(self.output_dir, f"{rid}_model_report.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        state._txt_report_path = path
        self._info(f"Text summary saved → {path}")
        return state

    # ───────────────────────── audit trail ─────────────────────────
    def _save_audit(self, state: PipelineState) -> PipelineState:
        path = os.path.join(self.output_dir, f"{state.run_id}_audit_trail.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(state.to_summary_dict(), f, indent=2, default=str)
        self._info(f"Audit trail saved → {path}")
        return state
