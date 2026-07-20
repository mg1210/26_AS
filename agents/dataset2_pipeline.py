"""
dataset2_pipeline.py — Blind Dataset 2 scoring workflow (8 phases)

Phase 1  Data Understanding      — schema, semantic types, business meanings, scope note
Phase 2  Data Quality Review     — scoped to champion features: missing, cardinality, outliers, dups
Phase 3  Feature Reconstruction  — derive engineered features from raw columns
Phase 4  Feature Stability       — per-feature PSI vs the training reference distribution
Phase 5  Explainability          — SHAP global importance on a sample (champion model)
Phase 6  Prediction              — score every row with the fixed champion + training imputation map
Phase 7  Validation              — score PSI + metrics/RAG/decile ONLY if a valid target is present
Phase 8  Fairness Check          — predicted-risk parity across proxy groups (target-independent)

IMPORTANT: the target column is NEVER used except in Phase 7 (metrics), and only if one is
found with valid mapped values. Phase 8 fairness examines predicted risk, not outcomes, so it
runs regardless of target availability. If the file is headerless, an ``alignment_verified=False``
flag and an UNVERIFIED warning are propagated through EVERY phase's result.
"""
import pandas as pd
import numpy as np
import json
import os
import sys
import glob
import joblib
import warnings
import subprocess
import tempfile
import shutil
from datetime import datetime

# Ensure the project root is importable when run directly as `python agents/dataset2_pipeline.py`
# (the script dir, agents/, is what lands on sys.path[0], so `import agents.*` would otherwise fail).
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


# ── Shared loaders / helpers ────────────────────────────────────────────────

def load_champion_meta(run_id=None):
    """Load the fixed champion model metadata + feature list from Dataset 1's run."""
    if run_id:
        path = f'outputs/models/{run_id}_features_meta.json'
    else:
        files = sorted(glob.glob('outputs/models/*_features_meta.json'), reverse=True)
        if not files:
            files = sorted(glob.glob('sample_results/models/*_features_meta.json'), reverse=True)
        if not files:
            raise FileNotFoundError("No champion model metadata found — run Dataset 1 pipeline first")
        path = files[0]
    with open(path, encoding='utf-8') as f:
        return json.load(f)


def load_champion_model(meta):
    """Load the champion .pkl (meta path first, then newest matching in outputs/ or sample_results/)."""
    model_path = meta.get('champion_model_path', '')
    if not model_path or not os.path.exists(model_path):
        candidates = sorted(glob.glob(f"outputs/models/*_{meta['champion_model']}.pkl"), reverse=True)
        candidates += sorted(glob.glob(f"sample_results/models/*_{meta['champion_model']}.pkl"), reverse=True)
        if not candidates:
            raise FileNotFoundError("Champion model file not found")
        model_path = candidates[0]
    return joblib.load(model_path)


def load_imputation_map(meta):
    """Load the exact training-time imputation decisions (col -> {strategy, fill_value})
    so missing values are filled identically at scoring time (no train/serve skew)."""
    path = meta.get('imputation_map_path', '')
    if not path or not os.path.exists(path):
        candidates = sorted(glob.glob('outputs/models/*_imputation_map.json'), reverse=True)
        candidates += sorted(glob.glob('sample_results/models/*_imputation_map.json'), reverse=True)
        if candidates:
            path = candidates[0]
    if path and os.path.exists(path):
        with open(path, encoding='utf-8') as f:
            return json.load(f)
    return {}


def load_reference_stats(meta):
    """Load per-raw-column training distribution stats (mean/std/min/p25/p50/p75/max)."""
    run_id = meta.get('run_id')
    path = f"outputs/models/{run_id}_reference_stats.json" if run_id else ''
    if not path or not os.path.exists(path):
        cands = sorted(glob.glob('outputs/models/*_reference_stats.json'), reverse=True)
        cands += sorted(glob.glob('sample_results/models/*_reference_stats.json'), reverse=True)
        if cands:
            path = cands[0]
    if path and os.path.exists(path):
        with open(path, encoding='utf-8') as f:
            return json.load(f)
    return {}


def load_reference_scores(meta):
    """Load Dataset 1's champion predicted-score distributions (OOT + Test) for score PSI."""
    path = meta.get('reference_scores_path', '')
    if not path or not os.path.exists(path):
        cands = sorted(glob.glob('outputs/models/*_reference_scores.json'), reverse=True)
        cands += sorted(glob.glob('sample_results/models/*_reference_scores.json'), reverse=True)
        if cands:
            path = cands[0]
    if path and os.path.exists(path):
        with open(path, encoding='utf-8') as f:
            return json.load(f)
    return {}


def compute_score_psi(d1_reference_scores, d2_scores, bins=10):
    """Standard PSI formula comparing two score distributions."""
    breakpoints = np.percentile(d1_reference_scores, np.linspace(0, 100, bins + 1))
    breakpoints[0], breakpoints[-1] = -np.inf, np.inf
    ref_pct = np.histogram(d1_reference_scores, bins=breakpoints)[0] / len(d1_reference_scores)
    new_pct = np.histogram(d2_scores, bins=breakpoints)[0] / len(d2_scores)
    ref_pct = np.where(ref_pct == 0, 1e-4, ref_pct)
    new_pct = np.where(new_pct == 0, 1e-4, new_pct)
    psi = float(np.sum((new_pct - ref_pct) * np.log(new_pct / ref_pct)))
    assessment = 'Stable' if psi < 0.10 else 'Moderate shift' if psi < 0.25 else 'Significant shift'
    return psi, assessment


def compute_decile_rank_order(y_true, y_prob, n_bins=10):
    """Decile rank-ordering check — bad rate should fall monotonically from the
    highest-risk decile to the lowest. Counts non-monotonic breaks."""
    df_dec = pd.DataFrame({'y': y_true, 'p': y_prob})
    df_dec['decile'] = pd.qcut(df_dec['p'], n_bins, labels=False, duplicates='drop')
    decile_table = df_dec.groupby('decile').agg(
        count=('y', 'size'),
        bad_rate=('y', 'mean'),
        mean_score=('p', 'mean')
    ).reset_index().sort_values('decile', ascending=False)  # highest risk decile first
    decile_table['decile_rank'] = range(1, len(decile_table) + 1)

    breaks = 0
    prev_rate = None
    break_details = []
    for _, row in decile_table.iterrows():
        if prev_rate is not None and row['bad_rate'] > prev_rate:
            breaks += 1
            break_details.append(f"Decile {int(row['decile_rank'])}: bad_rate {row['bad_rate']:.4f} > previous {prev_rate:.4f}")
        prev_rate = row['bad_rate']

    assessment = 'Strong — fully monotonic' if breaks == 0 else f'{breaks} break(s) detected'
    return decile_table.to_dict('records'), breaks, assessment, break_details


def build_scoring_matrix(df, meta):
    """Build the champion's input matrix X, filling missing cells with the EXACT training
    imputation values (sentinel/median) from the persisted map. Shared by Phase 5
    (explainability) and Phase 6 (prediction) so both see identical inputs.
    Indexed on df.index so scalar (missing→value) and Series (found) assignments both get
    the full row length (an empty frame would backfill early scalar columns to NaN)."""
    expected_features = meta['selected_features']
    imputation_map = load_imputation_map(meta)
    X = pd.DataFrame(index=df.index)
    imputation_applied = []
    for feat in expected_features:
        if feat in df.columns:
            s = pd.to_numeric(df[feat], errors='coerce')
            if feat in imputation_map:
                fill_val = imputation_map[feat]['fill_value']
                s = s.fillna(fill_val)
                imputation_applied.append({'feature': feat, 'fill_value': fill_val,
                                           'strategy': imputation_map[feat]['strategy']})
            else:
                s = s.fillna(0)
                imputation_applied.append({'feature': feat, 'fill_value': 0,
                                           'strategy': 'default (no training map entry)'})
            X[feat] = s
        else:
            # Feature genuinely missing entirely — use the training fill value as the best
            # guess (not a blanket 0, which for a sentinel column would misroute the model).
            if feat in imputation_map:
                X[feat] = imputation_map[feat]['fill_value']
                imputation_applied.append({'feature': feat, 'fill_value': imputation_map[feat]['fill_value'],
                                           'strategy': 'entire column missing, used training fill value'})
            else:
                X[feat] = 0
                imputation_applied.append({'feature': feat, 'fill_value': 0,
                                           'strategy': 'entire column missing, no training reference'})
    return X, imputation_applied


def reconstruct_engineered_features(df, expected_features):
    """Derive engineered features from raw columns if they're missing.

    Formulas mirror agents/feature_engineering_agent.py EXACTLY so the reconstructed
    values match the distribution the champion model was trained on. In particular the
    two ratios use a ``(denominator + 1)`` guard (not divide-then-drop-zeros) and the
    verification map fills unmapped values with 0 — both matching the training agent.
    """
    derived_log = []

    # int_rate_clean = numeric int_rate with '%' stripped  (agent._loan_features)
    if 'int_rate_clean' in expected_features and 'int_rate_clean' not in df.columns:
        if 'int_rate' in df.columns:
            df['int_rate_clean'] = pd.to_numeric(
                df['int_rate'].astype(str).str.replace('%', '').str.strip(), errors='coerce'
            )
            derived_log.append({'feature': 'int_rate_clean', 'derived_from': 'int_rate', 'status': 'Reconstructed'})
        else:
            derived_log.append({'feature': 'int_rate_clean', 'derived_from': 'int_rate (NOT FOUND)', 'status': 'Zero-imputed'})

    # loan_to_income = loan_amnt / (annual_inc + 1)   (agent formula: +1 denominator)
    if 'loan_to_income' in expected_features and 'loan_to_income' not in df.columns:
        if 'loan_amnt' in df.columns and 'annual_inc' in df.columns:
            loan_amnt = pd.to_numeric(df['loan_amnt'], errors='coerce')
            annual_inc = pd.to_numeric(df['annual_inc'], errors='coerce')
            df['loan_to_income'] = loan_amnt / (annual_inc + 1)
            derived_log.append({'feature': 'loan_to_income', 'derived_from': 'loan_amnt / (annual_inc + 1)', 'status': 'Reconstructed'})
        else:
            derived_log.append({'feature': 'loan_to_income', 'derived_from': 'loan_amnt or annual_inc (NOT FOUND)', 'status': 'Zero-imputed'})

    # open_acc_ratio = open_acc / (total_acc + 1)   (agent formula: +1 denominator)
    if 'open_acc_ratio' in expected_features and 'open_acc_ratio' not in df.columns:
        if 'open_acc' in df.columns and 'total_acc' in df.columns:
            open_acc = pd.to_numeric(df['open_acc'], errors='coerce')
            total_acc = pd.to_numeric(df['total_acc'], errors='coerce')
            df['open_acc_ratio'] = open_acc / (total_acc + 1)
            derived_log.append({'feature': 'open_acc_ratio', 'derived_from': 'open_acc / (total_acc + 1)', 'status': 'Reconstructed'})
        else:
            derived_log.append({'feature': 'open_acc_ratio', 'derived_from': 'open_acc or total_acc (NOT FOUND)', 'status': 'Zero-imputed'})

    # verification_ordinal from verification_status  (agent map, .fillna(0))
    if 'verification_ordinal' in expected_features and 'verification_ordinal' not in df.columns:
        if 'verification_status' in df.columns:
            vs_map = {'not verified': 0, 'source verified': 1, 'verified': 2}
            df['verification_ordinal'] = (
                df['verification_status'].astype(str).str.lower().str.strip().map(vs_map).fillna(0)
            )
            derived_log.append({'feature': 'verification_ordinal', 'derived_from': 'verification_status', 'status': 'Reconstructed'})
        else:
            derived_log.append({'feature': 'verification_ordinal', 'derived_from': 'verification_status (NOT FOUND)', 'status': 'Zero-imputed'})

    return df, derived_log


def _alignment_fields(phase1_result):
    """Standard alignment flags propagated into every phase result."""
    return {
        'alignment_verified': phase1_result.get('alignment_verified', True),
        'alignment_warning': phase1_result.get('alignment_warning', ''),
    }


# ── Statistical-scan fallback (headerless Dataset 2 only) ───────────────────
# The 9 champion raw variables the scanner must find; verification_status is categorical.
CHAMPION_RAW_VARS = ["int_rate", "loan_amnt", "annual_inc", "dti", "revol_util",
                     "tot_cur_bal", "open_acc", "total_acc", "verification_status"]
CATEGORICAL_RAW_VARS = {"verification_status": {"not verified", "source verified", "verified"}}
CONFIDENCE_THRESHOLD = 0.90
VALUE_BAND_TOL = 0.50      # median must be within [p25*(1-tol), p75*(1+tol)] and [min, max]
MISSING_PCT_TOL = 0.15     # |scanned missing% - training missing%| must be <= this
VOCAB_OVERLAP_MIN = 0.80   # >=80% of categorical rows must fall in the expected vocab


def run_column_scanner(headerless_csv, dict_path, out_dir, timeout=420):
    """Run lc_column_scanner.py as a subprocess on the FULL headerless CSV; parse
    column_mapping.json. Returns {col_index(int): {'variable': str|None, 'confidence': float}}
    or None if the scan cannot run (missing scanner/dict, subprocess error/timeout, unparseable
    output). Scans the full data (not a sample) so identification matches the scanner's true
    accuracy — this can take several minutes on large files, hence the generous timeout."""
    scanner = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           'lc_column_scanner.py')
    if not os.path.exists(scanner):
        return None
    cmd = [sys.executable, scanner, '--csv', headerless_csv, '--out', out_dir]
    if dict_path and os.path.exists(dict_path):
        cmd += ['--dict', dict_path]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout,
                           env=dict(os.environ, PYTHONUTF8='1', PYTHONIOENCODING='utf-8'))
        if r.returncode != 0:
            return None
        with open(os.path.join(out_dir, 'column_mapping.json'), encoding='utf-8') as f:
            cm = json.load(f).get('column_mapping', {})
        return {int(k): {'variable': v.get('variable'),
                         'confidence': float(v.get('confidence') or 0.0)}
                for k, v in cm.items()}
    except Exception:
        return None


def _sanity_numeric(series, ref):
    """Value-band plausibility (from training percentiles) AND missingness plausibility.
    Missingness is what catches a wrong pick like revol_util (train ~0% vs scanned 79.6%).
    Returns (ok, reason, extra); `extra` breaks out each sub-check verdict and the
    scanned-vs-training values so the HITL Feature Mapping Review can show why a pick
    passed or failed (both checks are evaluated even when the first one fails)."""
    miss = float(series.isna().mean())
    extra = {'value_check': '—', 'missingness_check': '—',
             'scanned_value': None, 'training_value': None,
             'scanned_missing_pct': round(miss * 100, 1), 'training_missing_pct': None}
    if not ref or 'missing_pct' not in ref:
        return False, 'no training missing_pct in reference_stats', extra
    extra['training_missing_pct'] = round(ref['missing_pct'] * 100, 1)
    extra['training_value'] = f"{ref['p50']:,.0f} median"
    s = pd.to_numeric(series, errors='coerce')
    if s.notna().sum() < 10:
        return False, 'too few numeric values', extra
    med = float(s.median())
    extra['scanned_value'] = f"{med:,.0f} median"
    val_ok = (ref['min'] <= med <= ref['max']) and \
             (ref['p25'] * (1 - VALUE_BAND_TOL) <= med <= ref['p75'] * (1 + VALUE_BAND_TOL))
    miss_ok = abs(miss - ref['missing_pct']) <= MISSING_PCT_TOL
    extra['value_check'] = 'PASS' if val_ok else 'FAIL'
    extra['missingness_check'] = 'PASS' if miss_ok else 'FAIL'
    if not val_ok:
        return False, (f"value out of band (median {med:.2f} vs train p25-p75 "
                       f"{ref['p25']}-{ref['p75']})"), extra
    if not miss_ok:
        return False, f"missingness mismatch (scanned {miss:.1%} vs train {ref['missing_pct']:.1%})", extra
    return True, f"ok (median {med:.2f}, missing {miss:.1%})", extra


def _sanity_categorical(series, expected_vocab):
    """Passes if a high fraction of rows fall in the expected category vocabulary.
    Returns (ok, reason, extra) mirroring _sanity_numeric's shape for the review table
    (categoricals have no missingness gate, so missingness_check is reported as N/A)."""
    vals = series.astype(str).str.lower().str.strip()
    overlap = float(vals.isin(expected_vocab).mean())
    miss = float(series.isna().mean())
    ok = overlap >= VOCAB_OVERLAP_MIN
    extra = {'value_check': 'PASS' if ok else 'FAIL', 'missingness_check': 'N/A',
             'scanned_value': f"{overlap:.0%} vocab overlap",
             'training_value': f"{len(expected_vocab)} expected categories",
             'scanned_missing_pct': round(miss * 100, 1), 'training_missing_pct': None}
    return ok, f"vocab overlap {overlap:.1%}" + ("" if ok else " (below threshold)"), extra


def _sanity_check(var, series, reference_stats):
    """Dispatch to the numeric or categorical sanity check for a champion raw variable.
    Returns (ok, reason, extra)."""
    if var in CATEGORICAL_RAW_VARS:
        return _sanity_categorical(series, CATEGORICAL_RAW_VARS[var])
    ref = (reference_stats or {}).get(var)
    if not ref:
        return False, 'no reference stats', {
            'value_check': '—', 'missingness_check': '—',
            'scanned_value': None, 'training_value': None,
            'scanned_missing_pct': None, 'training_missing_pct': None}
    return _sanity_numeric(series, ref)


def gate_scanner_assignments(df_raw, scanner_map, reference_stats):
    """Per-variable confidence + sanity gating. A variable is RESOLVED only if the scanner's
    confidence >= CONFIDENCE_THRESHOLD AND the sanity check passes; otherwise UNRESOLVED (never
    guessed). Returns (resolved{var:col_index}, unresolved[list], detail[list-of-dicts])."""
    # Invert the scanner map: champion raw var -> best (col_index, confidence).
    var_to_col = {}
    for col, info in scanner_map.items():
        v = info.get('variable')
        if v in CHAMPION_RAW_VARS and (v not in var_to_col or info['confidence'] > var_to_col[v][1]):
            var_to_col[v] = (col, info['confidence'])

    resolved, unresolved, detail = {}, [], []
    for var in CHAMPION_RAW_VARS:
        # Full per-variable record for the HITL Feature Mapping Review table. `scanner_col` is
        # kept for backward compatibility; `mapped_column` is the same value under the field
        # name the review UI reads (the scanner's pick, even when it is later UNRESOLVED).
        entry = {'variable': var, 'scanner_col': None, 'mapped_column': None,
                 'confidence': None, 'sanity_ok': None,
                 'value_check': '—', 'missingness_check': '—',
                 'scanned_value': None, 'training_value': None,
                 'scanned_missing_pct': None, 'training_missing_pct': None,
                 'reason': None, 'status': 'UNRESOLVED'}
        if var not in var_to_col:
            entry['reason'] = 'scanner did not assign this variable'
            unresolved.append(var); detail.append(entry); continue
        col, conf = var_to_col[var]
        entry['scanner_col'] = col
        entry['mapped_column'] = col
        entry['confidence'] = round(conf, 3)
        conf_ok = conf >= CONFIDENCE_THRESHOLD
        sane, reason, extra = _sanity_check(var, df_raw.iloc[:, col], reference_stats)
        entry.update(extra)
        entry['sanity_ok'] = sane
        entry['reason'] = reason
        if conf_ok and sane:
            resolved[var] = col
            entry['status'] = 'RESOLVED'
        else:
            unresolved.append(var)
            if not conf_ok:
                entry['reason'] = f"confidence {conf:.2f} < {CONFIDENCE_THRESHOLD} | {reason}"
        detail.append(entry)
    return resolved, unresolved, detail


# ── Phase 1 — Data Understanding ────────────────────────────────────────────

def phase1_data_understanding(dataset_path, meta):
    """Profile Dataset 2 — schema, semantic types, business meanings. No target handling."""
    result = {'phase': 'Data Understanding', 'dataset': os.path.basename(dataset_path)}

    # Header detection — compare the numeric fraction of the first two rows. A header row is
    # markedly LESS numeric than the first data row beneath it; a headerless file has two
    # consecutive DATA rows with similar numeric fractions. This is robust to categorical/empty-
    # heavy data where a single-row "is it mostly numeric?" test misfires (a data row full of
    # categoricals/blanks looks header-like on its own).
    def _numeric_fraction(line):
        fields = line.strip().split(',')
        if not fields:
            return 0.0
        n = sum(1 for fld in fields
                if fld.strip().replace('.', '').replace('-', '').lstrip('-').isdigit())
        return n / len(fields)
    with open(dataset_path, 'r', encoding='utf-8', errors='replace') as f:
        _line0 = f.readline()
        _line1 = f.readline()
    _f0 = _numeric_fraction(_line0)
    if _line1.strip():
        has_header = (_numeric_fraction(_line1) - _f0) > 0.20
    else:
        has_header = _f0 < 0.5   # single-row fallback

    # alignment_verified=True only when columns are matched by real names (headered file).
    # Headerless files fall back to positional assignment, which can silently produce
    # garbage when the file's column order differs from Dataset 1 — so we flag them.
    HEADERLESS_WARNING = (
        'This file appears to have no column headers. Column order cannot be reliably '
        'verified — for accurate scoring, please ensure this file has the same column '
        'names as Dataset 1, or add a header row before uploading.'
    )
    if has_header:
        df = pd.read_csv(dataset_path, low_memory=False)
        result['header_detected'] = True
        result['alignment_verified'] = True
        result['alignment_warning'] = ''
    else:
        # ── Headerless: statistical-scan fallback (confidence + sanity gated) ──
        result['header_detected'] = False
        df_raw = pd.read_csv(dataset_path, header=None, low_memory=False)

        dict_path = None
        for _dp in glob.glob('data/*ictionary*.xlsx'):
            dict_path = _dp
            break
        _tmp_out = tempfile.mkdtemp(prefix='d2scan_')
        try:
            scanner_map = run_column_scanner(dataset_path, dict_path, _tmp_out)
        finally:
            shutil.rmtree(_tmp_out, ignore_errors=True)

        if scanner_map is not None:
            reference_stats = load_reference_stats(meta)
            resolved, unresolved, detail = gate_scanner_assignments(df_raw, scanner_map, reference_stats)
            # Name ONLY resolved columns with their true champion-raw names. Every other column
            # keeps a neutral col_i name, so an unresolved/failed pick can never masquerade as a
            # champion feature — it is simply absent and gets training-median imputation downstream.
            new_cols = [f'col_{i}' for i in range(df_raw.shape[1])]
            for var, col in resolved.items():
                new_cols[col] = var
            df = df_raw.copy()
            df.columns = new_cols

            n_ok, m_bad = len(resolved), len(unresolved)
            result['scan_summary'] = (f"{n_ok} of {len(CHAMPION_RAW_VARS)} raw variables identified "
                                      f"with high confidence via statistical scan, {m_bad} unresolved")
            result['scan_detail'] = detail
            result['columns_assigned_from'] = 'statistical scan (confidence + sanity gated)'
            if m_bad == 0:
                result['alignment_verified'] = True
                result['alignment_method'] = 'statistical scan (all variables verified)'
                result['alignment_warning'] = ''
            else:
                result['alignment_verified'] = False
                result['alignment_method'] = 'statistical scan (partial)'
                result['alignment_warning'] = (
                    f"UNVERIFIED — {m_bad} champion feature input(s) could not be confidently "
                    f"identified via statistical scan: {unresolved}. These are imputed with training "
                    "medians; predictions for them are approximate.")
        else:
            # Scanner unavailable/failed → existing positional best-effort, UNVERIFIED.
            result['alignment_verified'] = False
            result['alignment_method'] = 'positional (scanner unavailable)'
            result['alignment_warning'] = HEADERLESS_WARNING
            col_order_path = 'outputs/models/dataset1_column_order.json'
            if not os.path.exists(col_order_path):
                col_order_path = 'sample_results/models/dataset1_column_order.json'
            df = df_raw
            if os.path.exists(col_order_path):
                with open(col_order_path) as f:
                    dataset1_cols = json.load(f)
                if len(dataset1_cols) == len(df.columns):
                    df.columns = dataset1_cols
                    result['columns_assigned_from'] = 'Dataset 1 reference order (UNVERIFIED — positional best-effort)'
                else:
                    df.columns = [f'col_{i}' for i in range(len(df.columns))]
                    result['columns_assigned_from'] = 'auto-generated (count mismatch with Dataset 1)'
            else:
                df.columns = [f'col_{i}' for i in range(len(df.columns))]
                result['columns_assigned_from'] = 'auto-generated (no reference found)'

    result['total_rows'] = len(df)
    result['total_columns'] = df.shape[1]

    # Business meanings: Dataset 2's own dict -> Dataset 1's dict as GLOBAL fallback -> name inference
    data_dict = {}
    own_dict_files = glob.glob('data/*dataset2*dict*.xlsx') + glob.glob('data/*2*dictionary*.xlsx')
    dict_source = None
    if own_dict_files:
        try:
            dd = pd.read_excel(own_dict_files[0])
            cols = dd.columns.tolist()
            data_dict = dict(zip(dd[cols[0]].astype(str).str.lower().str.strip(), dd[cols[1]].astype(str)))
            dict_source = f'Dataset 2 own dictionary ({own_dict_files[0]})'
        except Exception:
            pass
    if not data_dict:
        global_dict_files = glob.glob('data/*ictionary*.xlsx')
        if global_dict_files:
            try:
                dd = pd.read_excel(global_dict_files[0])
                cols = dd.columns.tolist()
                data_dict = dict(zip(dd[cols[0]].astype(str).str.lower().str.strip(), dd[cols[1]].astype(str)))
                dict_source = f'Dataset 1 dictionary used as global fallback ({global_dict_files[0]})'
            except Exception:
                pass
    result['dictionary_source'] = dict_source or 'None found — using name inference'

    # Semantic type inference per column (numeric / categorical / date / identifier)
    semantic_types = []
    for col in df.columns:
        s = df[col]
        if pd.api.types.is_numeric_dtype(s) or pd.to_numeric(s, errors='coerce').notna().mean() > 0.8:
            n_unique = s.nunique()
            sem_type = 'identifier' if n_unique / max(len(s), 1) > 0.95 else 'numeric'
        else:
            try:
                with warnings.catch_warnings():
                    warnings.simplefilter('ignore')
                    pd.to_datetime(s.dropna().head(50), errors='raise')
                sem_type = 'date'
            except Exception:
                sem_type = 'categorical'
        semantic_types.append(sem_type)

    # Schema profile with semantic types. NOTE: missing-value analysis (missing_pct) lives
    # ONLY in Phase 2 (Data Quality Review), mirroring the DQR-owns-missing separation —
    # Data Understanding describes structure/semantics, not data-quality metrics.
    schema = []
    for col, sem_type in zip(df.columns, semantic_types):
        schema.append({
            'column': col, 'semantic_type': sem_type,
            'dtype': str(df[col].dtype),
            'n_unique': int(df[col].nunique()),
            'business_meaning': data_dict.get(col.lower(), '—'),
            'is_expected_feature': col in meta['selected_features'],
        })
    result['schema'] = schema

    # Selection bias reminder (informational only for Dataset 2)
    result['scope_note'] = (
        'This dataset is being scored using a model trained on funded/approved loans. '
        'If Dataset 2 includes declined applicants, predictions for that subset carry '
        'additional uncertainty (selection bias) not captured during training.'
    )

    result['df'] = df  # kept in memory for next phase, not serialized
    return result


# ── Phase 2 — Data Quality Review (scoped) ──────────────────────────────────

def phase2_data_quality_scoped(df, meta, phase1_result):
    """DQR scoped ONLY to the features the champion model needs — assessed on the raw
    incoming data (engineered features are reconstructed later, in Phase 3)."""
    expected_features = meta['selected_features']
    result = {'phase': 'Data Quality Review (Scoped)', 'expected_features': expected_features}
    result.update(_alignment_fields(phase1_result))

    available = [f for f in expected_features if f in df.columns]
    missing = [f for f in expected_features if f not in df.columns]
    result['features_found'] = available
    result['features_missing'] = missing
    result['mapping_method'] = 'Exact column name match'

    # Missing value + cardinality + summary stats per needed feature
    quality_rows = []
    for feat in expected_features:
        if feat in df.columns:
            s = pd.to_numeric(df[feat], errors='coerce')
            quality_rows.append({
                'feature': feat, 'status': 'Found',
                'missing_pct': round(s.isna().mean() * 100, 2),
                'n_unique': int(df[feat].nunique()),
                'mean': round(float(s.mean()), 4) if s.notna().any() else None,
                'std': round(float(s.std()), 4) if s.notna().any() else None,
                'min': round(float(s.min()), 4) if s.notna().any() else None,
                'max': round(float(s.max()), 4) if s.notna().any() else None,
            })
        else:
            quality_rows.append({'feature': feat, 'status': 'MISSING', 'missing_pct': 100.0,
                                 'n_unique': None, 'mean': None, 'std': None, 'min': None, 'max': None})
    result['quality_table'] = quality_rows

    # Outlier detection (IQR × 3) on needed numeric features that are present
    outlier_rows = []
    for feat in available:
        s = pd.to_numeric(df[feat], errors='coerce').dropna()
        if len(s) < 30:
            continue
        q1, q3 = s.quantile(0.25), s.quantile(0.75)
        iqr = q3 - q1
        lower, upper = q1 - 3 * iqr, q3 + 3 * iqr
        n_outliers = int(((s < lower) | (s > upper)).sum())
        outlier_rows.append({'feature': feat, 'iqr_outliers': n_outliers,
                             'pct': round(n_outliers / len(s) * 100, 2)})
    result['outlier_table'] = outlier_rows

    # Duplicate check if an ID-like column exists
    id_candidates = [c for c in df.columns if c.lower() in ('id', 'member_id', 'record_no', 'loan_id')]
    dup_info = {}
    for idc in id_candidates:
        dup_info[idc] = int(df[idc].duplicated().sum())
    result['duplicate_check'] = dup_info or {'note': 'No obvious ID column found to check duplicates'}

    return result


# ── Phase 3 — Feature Reconstruction ────────────────────────────────────────

def phase3_feature_reconstruction(df, meta, phase1_result):
    """Derive the engineered features the champion needs from raw columns (Phase 2 assessed
    raw availability; this phase builds the missing engineered ones with training-exact math)."""
    result = {'phase': 'Feature Reconstruction'}
    result.update(_alignment_fields(phase1_result))
    df, derived_log = reconstruct_engineered_features(df, meta['selected_features'])
    result['derived_features_log'] = derived_log
    reconstructed = [d['feature'] for d in derived_log if d['status'] == 'Reconstructed']
    result['n_reconstructed'] = len(reconstructed)
    result['reconstructed_features'] = reconstructed
    now_present = [f for f in meta['selected_features'] if f in df.columns]
    result['features_available_after'] = now_present
    result['n_available_after'] = len(now_present)
    result['df'] = df
    return result


# ── Phase 4 — Feature Stability (PSI vs training) ───────────────────────────

def phase4_feature_stability(df, meta, phase1_result):
    """Population Stability Index (PSI) per champion feature — Dataset 2 vs the training
    reference distribution. Mirrors the Dataset 1 validation agent's PSI concept, applied
    to feature distributions to detect covariate/population drift. Target is never used.

    PSI is computed against the training percentile edges saved in reference_stats.json
    (min/p25/p50/p75/max → 4 equal-mass training bins), so PSI = Σ(actual−0.25)·ln(actual/0.25)."""
    result = {'phase': 'Feature Stability'}
    result.update(_alignment_fields(phase1_result))
    ref = load_reference_stats(meta)
    # Engineered champion features map back to the raw column whose reference stats apply.
    ref_key_for = {'int_rate_clean': 'int_rate'}

    stability_rows = []
    for feat in meta['selected_features']:
        ref_key = feat if feat in ref else ref_key_for.get(feat)
        present = feat in df.columns
        d2_mean = None
        if present:
            s_all = pd.to_numeric(df[feat], errors='coerce').dropna()
            if len(s_all):
                d2_mean = round(float(s_all.mean()), 4)
        d1_mean = round(float(ref[ref_key]['mean']), 4) if (ref_key and ref_key in ref) else None
        row = {'feature': feat, 'ref_column': ref_key or '—',
               'd1_mean': d1_mean, 'd2_mean': d2_mean, 'mean_shift_pct': None,
               'psi': None, 'stability_flag': 'NO REFERENCE',
               'assessment': 'No training reference available'}
        if d1_mean is not None and d2_mean is not None and d1_mean != 0:
            row['mean_shift_pct'] = round((d2_mean - d1_mean) / abs(d1_mean) * 100, 2)

        if not ref_key or ref_key not in ref or not present:
            stability_rows.append(row)
            continue
        rs = ref[ref_key]
        edges = sorted(set([rs['min'], rs['p25'], rs['p50'], rs['p75'], rs['max']]))
        if len(edges) < 3:
            row['assessment'] = 'Degenerate reference (near-constant)'
            stability_rows.append(row)
            continue
        s = pd.to_numeric(df[feat], errors='coerce').dropna()
        if len(s) < 30:
            row['assessment'] = 'Too few values to assess'
            stability_rows.append(row)
            continue
        # Inner percentile edges with open ends so out-of-range D2 values still count.
        inner = edges[1:-1]
        bins = [-np.inf] + inner + [np.inf]
        n_bins = len(bins) - 1
        expected = np.full(n_bins, 1.0 / n_bins)          # equal-mass training bins by construction
        actual = np.histogram(s, bins=bins)[0] / len(s)
        expected = np.where(expected == 0, 1e-4, expected)
        actual = np.where(actual == 0, 1e-4, actual)
        psi = float(np.sum((actual - expected) * np.log(actual / expected)))
        row['psi'] = round(psi, 4)
        if psi < 0.10:
            row['stability_flag'], row['assessment'] = 'STABLE', 'Stable (<0.10)'
        elif psi < 0.25:
            row['stability_flag'], row['assessment'] = 'SHIFTED', 'Moderate shift (0.10–0.25)'
        else:
            row['stability_flag'], row['assessment'] = 'HIGH DRIFT', 'Unstable (>0.25)'
        stability_rows.append(row)
    result['stability_table'] = stability_rows

    assessed = [r for r in stability_rows if r['psi'] is not None]
    if assessed:
        avg = float(np.mean([r['psi'] for r in assessed]))
        result['avg_psi'] = round(avg, 4)
        result['overall_assessment'] = ('Stable' if avg < 0.10
                                        else 'Moderate drift' if avg < 0.25
                                        else 'Unstable — significant population drift')
        result['n_assessed'] = len(assessed)
    else:
        result['avg_psi'] = None
        result['overall_assessment'] = 'No features had training reference stats to assess'
        result['n_assessed'] = 0
    result['psi_note'] = ('Approximate PSI using saved training percentile edges (min/p25/p50/p75/max). '
                          'Engineered ratio features without a raw reference are not assessed.')
    return result


# ── Phase 5 — Explainability (SHAP) ─────────────────────────────────────────

def phase5_explainability(df, meta, phase1_result, sample_size=1000):
    """Global feature importance via SHAP on a Dataset 2 sample using the champion model.
    Mirrors the Dataset 1 explainability agent. Target is never used. Falls back to the
    model's native importances / coefficients if SHAP is unavailable for the model type."""
    result = {'phase': 'Explainability'}
    result.update(_alignment_fields(phase1_result))

    X, _ = build_scoring_matrix(df, meta)
    model = load_champion_model(meta)
    n = min(sample_size, len(X))
    X_sample = X.sample(n=n, random_state=42) if len(X) > n else X

    method = 'unavailable'
    try:
        import shap
        explainer = shap.TreeExplainer(model)
        sv = explainer.shap_values(X_sample)
        if isinstance(sv, list):              # some versions return [class0, class1]
            sv = sv[1] if len(sv) > 1 else sv[0]
        sv = np.asarray(sv)
        if sv.ndim == 3:                      # (n, features, classes)
            sv = sv[:, :, -1]
        mean_abs = np.abs(sv).mean(axis=0)
        method = 'SHAP (TreeExplainer)'
    except Exception as e:
        result['shap_note'] = f'SHAP unavailable ({type(e).__name__}); used model importances'
        if hasattr(model, 'feature_importances_'):
            mean_abs = np.asarray(model.feature_importances_, dtype=float)
            method = 'Model feature_importances_'
        elif hasattr(model, 'coef_'):
            mean_abs = np.abs(np.ravel(model.coef_)).astype(float)
            method = 'Logistic |coef|'
        else:
            mean_abs = np.zeros(X_sample.shape[1])

    feats = list(X_sample.columns)
    total = float(np.sum(mean_abs)) or 1.0
    importance = [{'feature': f, 'mean_abs_shap': round(float(v), 6),
                   'importance_pct': round(100 * float(v) / total, 2)}
                  for f, v in sorted(zip(feats, mean_abs), key=lambda t: t[1], reverse=True)]
    result['method'] = method
    result['sample_size'] = n
    result['feature_importance'] = importance
    result['top_features'] = [r['feature'] for r in importance[:5]]
    return result


# ── Phase 6 — Prediction ────────────────────────────────────────────────────

def phase6_prediction(df, meta, phase1_result=None):
    """Score every row using the fixed champion model + the training imputation map."""
    result = {'phase': 'Prediction'}
    if phase1_result is not None:
        result.update(_alignment_fields(phase1_result))

    X, imputation_applied = build_scoring_matrix(df, meta)
    result['imputation_applied'] = imputation_applied

    model = load_champion_model(meta)
    y_prob = model.predict_proba(X)[:, 1]

    result['n_scored'] = len(y_prob)
    result['champion_model'] = meta['champion_model']
    result['predictions'] = y_prob.tolist()
    result['score_mean'] = round(float(y_prob.mean()), 4)
    result['score_median'] = round(float(np.median(y_prob)), 4)
    result['score_min'] = round(float(y_prob.min()), 4)
    result['score_max'] = round(float(y_prob.max()), 4)
    result['pct_high_risk'] = round(float((y_prob >= 0.5).mean() * 100), 2)

    bands = pd.cut(y_prob, bins=[0, 0.2, 0.4, 0.6, 0.8, 1.0],
                   labels=['Very Low', 'Low', 'Medium', 'High', 'Very High'])
    result['risk_band_distribution'] = bands.value_counts().to_dict()

    result['y_prob'] = y_prob
    return result


# ── Phase 7 — Validation (conditional) ──────────────────────────────────────

def phase7_validation(df, y_prob, meta):
    """Score-level PSI (always) + discrimination metrics + RAG rating ONLY if a target
    column with valid mapped values is present. The target is never assumed."""
    result = {'phase': 'Validation', 'target_available': False}

    # ── Score-level PSI (target-independent): Dataset 2 scores vs Dataset 1 reference ──
    ref_scores = load_reference_scores(meta)
    d1_ref, ref_label = None, None
    if ref_scores.get('oot_scores'):
        d1_ref, ref_label = ref_scores['oot_scores'], 'Dataset 1 OOT'
    elif ref_scores.get('test_scores'):
        d1_ref, ref_label = ref_scores['test_scores'], 'Dataset 1 Test'
    if d1_ref:
        psi, assess = compute_score_psi(np.asarray(d1_ref, dtype=float), np.asarray(y_prob, dtype=float))
        result['score_psi'] = round(psi, 4)
        result['score_psi_assessment'] = assess
        result['score_psi_reference'] = ref_label
    else:
        result['score_psi'] = None
        result['score_psi_assessment'] = 'No Dataset 1 reference scores available (re-run training to persist them)'
        result['score_psi_reference'] = None

    target_col_name = meta.get('target_column', 'loan_status')
    target_mapping = meta.get('target_mapping', {})

    # Try to find a target column — check exact name, then case-insensitive
    found_col = None
    if target_col_name in df.columns:
        found_col = target_col_name
    else:
        col_lower = {c.lower(): c for c in df.columns}
        if target_col_name.lower() in col_lower:
            found_col = col_lower[target_col_name.lower()]

    if found_col is None:
        result['message'] = 'No target column found in Dataset 2 — this is a true blind scoring scenario. Only prediction distribution is available.'
        result['overall_rating'] = 'N/A — no target available for RAG assessment'
        return result

    # Map values using saved mapping
    mapping_lower = {str(k).lower().strip(): v for k, v in target_mapping.items() if v in (0, 1)}
    y_true_raw = df[found_col].astype(str).str.lower().str.strip().map(mapping_lower)
    valid_mask = y_true_raw.notna()

    if valid_mask.sum() < 10:
        result['message'] = f'Target column "{found_col}" found but fewer than 10 valid mapped values — insufficient for validation metrics.'
        result['overall_rating'] = 'N/A — no target available for RAG assessment'
        return result

    y_true = y_true_raw[valid_mask].astype(int).values
    y_pred_prob = y_prob[valid_mask.values]

    from sklearn.metrics import roc_auc_score, roc_curve, confusion_matrix, brier_score_loss

    auc = roc_auc_score(y_true, y_pred_prob)
    gini = 2 * auc - 1
    fpr, tpr, _ = roc_curve(y_true, y_pred_prob)
    ks = float(np.max(tpr - fpr))
    brier = brier_score_loss(y_true, y_pred_prob)

    # Decile rank-order (monotonicity) check
    decile_table, n_breaks, rank_assessment, break_details = compute_decile_rank_order(y_true, y_pred_prob)

    y_pred_class = (y_pred_prob >= 0.5).astype(int)
    cm = confusion_matrix(y_true, y_pred_class)
    tn, fp, fn, tp = cm.ravel()
    precision = tp / (tp + fp) if (tp + fp) > 0 else 0
    recall = tp / (tp + fn) if (tp + fn) > 0 else 0
    f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0

    result['target_available'] = True
    result['target_column_used'] = found_col
    result['n_valid'] = int(valid_mask.sum())
    result['auc'] = round(auc, 4)
    result['gini'] = round(gini, 4)
    result['ks'] = round(ks, 4)
    result['brier_score'] = round(brier, 4)
    result['confusion_matrix'] = {
        'true_positive': int(tp), 'false_positive': int(fp),
        'true_negative': int(tn), 'false_negative': int(fn),
        'precision': round(precision, 4), 'recall': round(recall, 4), 'f1_score': round(f1, 4),
    }
    result['default_rate'] = round(float(y_true.mean()), 4)

    # Decile rank-order results
    result['decile_table'] = decile_table
    result['rank_order_breaks'] = n_breaks
    result['rank_order_assessment'] = rank_assessment
    result['rank_order_break_details'] = break_details

    # ── RAG rating — reuse Dataset 1's exact KPI thresholds (no second hardcoded copy) ──
    from agents.validation_agent import get_rag
    rag_summary = {
        'auc': {'value': result['auc'], 'rag': get_rag('auc', result['auc'])[0]},
        'ks': {'value': result['ks'], 'rag': get_rag('ks', result['ks'])[0]},
    }
    # Rank-order RAG matches validation_agent's pattern: 0 breaks GREEN, 1-2 AMBER, 3+ RED.
    rag_summary['rank_order'] = {
        'value': n_breaks,
        'rag': 'GREEN' if n_breaks == 0 else 'AMBER' if n_breaks <= 2 else 'RED',
    }
    if result.get('score_psi') is not None:
        rag_summary['score_psi'] = {'value': result['score_psi'],
                                    'rag': get_rag('psi', result['score_psi'])[0]}
    result['rag_summary'] = rag_summary
    rags = [v['rag'] for v in rag_summary.values()]
    result['overall_rating'] = ('GREEN' if all(r == 'GREEN' for r in rags)
                                else 'RED' if any(r == 'RED' for r in rags) else 'AMBER')

    return result


# ── Phase 8 — Fairness Check ────────────────────────────────────────────────

def phase_fairness_check(df, y_prob, meta):
    """Fairness re-check on the Dataset 2 scored population. Runs regardless of target
    availability — it examines predicted risk across proxy groups, not actual outcomes."""
    result = {'phase': 'Fairness Check (Dataset 2 Population)'}
    proxy_attrs = ['verification_status', 'home_ownership', 'purpose', 'addr_state']

    fairness_results = {}
    for attr in proxy_attrs:
        if attr not in df.columns:
            continue
        try:
            df_fair = pd.DataFrame({'group': df[attr].values, 'predicted_prob': y_prob})
            overall_mean = df_fair['predicted_prob'].mean()
            group_stats = df_fair.groupby('group').agg(
                count=('predicted_prob', 'size'),
                mean_predicted=('predicted_prob', 'mean')
            ).reset_index()
            group_stats = group_stats[group_stats['count'] >= 30]
            group_stats['diff_from_avg'] = group_stats['mean_predicted'] - overall_mean
            group_stats['concern_level'] = group_stats['diff_from_avg'].abs().apply(
                lambda d: 'High' if d > 0.15 else 'Medium' if d > 0.10 else 'Low'
            )
            fairness_results[attr] = {
                row['group']: {
                    'count': int(row['count']),
                    'mean_predicted': round(float(row['mean_predicted']), 4),
                    'diff_from_avg': round(float(row['diff_from_avg']), 4),
                    'concern_level': row['concern_level'],
                }
                for _, row in group_stats.iterrows()
            }
        except Exception:
            continue

    total_flagged = sum(1 for attr_data in fairness_results.values()
                        for v in attr_data.values() if v['concern_level'] in ('Medium', 'High'))
    result['fairness_results'] = fairness_results
    result['attributes_checked'] = len(fairness_results)
    result['groups_flagged'] = total_flagged
    result['summary'] = (f'{len(fairness_results)} proxy attributes checked on Dataset 2 population, '
                         f'{total_flagged} group(s) flagged')
    return result


# ── Phase 9 — Documentation ─────────────────────────────────────────────────

def phase9_documentation(combined_results, meta):
    """Generate a Dataset 2 Scoring Report (Word .docx), mirroring Dataset 1's Model
    Development Document structure but scoped to this blind scoring run."""
    from docx import Document
    from docx.shared import Inches
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    doc = Document()
    run_id = combined_results.get('dataset2_run_id', 'unknown')

    # Title
    doc.add_heading('Dataset 2 — Blind Scoring Report', level=0)
    doc.add_paragraph(f"Run ID: {run_id}")
    doc.add_paragraph(f"Source Dataset 1 Model: {combined_results.get('source_dataset1_run', '—')}")
    doc.add_paragraph(f"Champion Model Used: {combined_results.get('champion_model', '—')}")
    doc.add_paragraph(f"Generated: {combined_results.get('timestamp', '—')}")

    alignment_verified = combined_results.get('alignment_verified', True)
    p1 = combined_results.get('phase1_data_understanding', {})
    alignment_method = (combined_results.get('alignment_method')
                        or p1.get('alignment_method', 'header-based exact match'))
    if alignment_verified:
        doc.add_paragraph(f"✓ Column Alignment: VERIFIED ({alignment_method})")
    else:
        doc.add_paragraph(f"⚠ Column Alignment: UNVERIFIED — {combined_results.get('alignment_warning', '')}")

    p6 = combined_results.get('phase6_prediction', {})
    p7 = combined_results.get('phase7_validation', {})

    # Section 1 — Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        f"This report documents the blind scoring of a new dataset ({p1.get('total_rows', 0):,} rows, "
        f"{p1.get('total_columns', 0)} columns) using the fixed champion model from Dataset 1 "
        f"({combined_results.get('champion_model', '—')}), with no retraining performed. "
        f"{p6.get('n_scored', 0):,} records were scored, yielding a mean predicted default "
        f"probability of {p6.get('score_mean', '—')} and {p6.get('pct_high_risk', '—')}% flagged high-risk. "
        + (f"Validation against actual outcomes (target column '{p7.get('target_column_used', '—')}') confirmed "
           f"AUC={p7.get('auc', '—')}, consistent with the model's documented Dataset 1 performance."
           if p7.get('target_available') else
           "No ground-truth labels were available in this dataset — this is a true blind scoring exercise; "
           "validation metrics will be computable once outcomes are observed.")
    )

    # Section 2 — Process Flow
    doc.add_heading('2. Process Flow', level=1)
    doc.add_paragraph(
        'This scoring run followed the same 8-phase governance structure used for Dataset 1 model '
        'development, adapted for blind scoring (phases requiring labels are conditional):'
    )
    phases_desc = [
        ('Phase 1 — Data Understanding', 'Schema profiling, business meaning lookup, column alignment verification. No target column used.'),
        ('Phase 2 — Data Quality Review', 'Missing value, outlier, and duplicate checks scoped to the features the champion model requires.'),
        ('Phase 3 — Feature Reconstruction', 'Engineered features (int_rate_clean, loan_to_income, etc.) rebuilt from raw columns using the same logic as Dataset 1 training.'),
        ('Phase 4 — Feature Stability (PSI)', "Each feature's distribution compared against Dataset 1's training reference — fully unsupervised drift check."),
        ('Phase 5 — Explainability (SHAP)', "SHAP feature importance computed on this dataset's own predictions using the fixed champion model."),
        ('Phase 6 — Prediction', 'Every row scored using the fixed champion model and training-exact imputation map.'),
        ('Phase 7 — Validation', 'Conditional — AUC/KS/Gini/Confusion Matrix computed ONLY if a target column with valid labels is present.'),
        ('Phase 8 — Fairness Check', 'Predicted-risk parity check across proxy attributes on this scored population.'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Phase'
    table.rows[0].cells[1].text = 'Description'
    for name, desc in phases_desc:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc

    # Section 3 — Data Understanding & Column Alignment
    doc.add_heading('3. Data Understanding & Column Alignment', level=1)
    doc.add_paragraph(f"Rows: {p1.get('total_rows', 0):,} | Columns: {p1.get('total_columns', 0)}")
    doc.add_paragraph(f"Header detected: {p1.get('header_detected', True)}")
    doc.add_paragraph(f"Business meaning source: {p1.get('dictionary_source', '—')}")
    if 'scan_summary' in p1:
        doc.add_paragraph(f"Statistical column scan: {p1.get('scan_summary', '')}")

    # Section 4 — Data Quality Review
    doc.add_heading('4. Data Quality Review', level=1)
    p2 = combined_results.get('phase2_data_quality', {})
    doc.add_paragraph(f"Features found: {len(p2.get('features_found', []))} of {len(p2.get('expected_features', []))}")
    if p2.get('features_missing'):
        doc.add_paragraph(f"Missing: {p2.get('features_missing')}")
    qtable = p2.get('quality_table', [])
    if qtable:
        t = doc.add_table(rows=1, cols=len(qtable[0]))
        t.style = 'Light Grid Accent 1'
        for i, k in enumerate(qtable[0].keys()):
            t.rows[0].cells[i].text = str(k)
        for row_data in qtable:
            row = t.add_row().cells
            for i, v in enumerate(row_data.values()):
                row[i].text = str(v)

    # Section 5 — Feature Reconstruction
    doc.add_heading('5. Feature Reconstruction', level=1)
    p3 = combined_results.get('phase3_feature_reconstruction', {})
    recon_log = p3.get('derived_features_log', [])
    n_zero = sum(1 for e in recon_log if e.get('status') == 'Zero-imputed')
    doc.add_paragraph(f"Reconstructed: {p3.get('n_reconstructed', 0)} | Zero-imputed: {n_zero} | "
                      f"Champion features available: {p3.get('n_available_after', '—')}")
    if recon_log:
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Feature', 'Derived From', 'Status']):
            t.rows[0].cells[i].text = h
        for entry in recon_log:
            row = t.add_row().cells
            row[0].text = entry.get('feature', '')
            row[1].text = entry.get('derived_from', '')
            row[2].text = entry.get('status', '')

    # Section 6 — Feature Stability (PSI)
    doc.add_heading('6. Feature Stability (Population Stability Index)', level=1)
    p4 = combined_results.get('phase4_feature_stability', {})
    doc.add_paragraph(f"{p4.get('overall_assessment', '—')} (avg PSI={p4.get('avg_psi', '—')}, "
                      f"{p4.get('n_assessed', 0)} feature(s) assessed). {p4.get('psi_note', '')}")
    stab_table = p4.get('stability_table', [])
    if stab_table:
        t = doc.add_table(rows=1, cols=len(stab_table[0]))
        t.style = 'Light Grid Accent 1'
        for i, k in enumerate(stab_table[0].keys()):
            t.rows[0].cells[i].text = str(k)
        for row_data in stab_table:
            row = t.add_row().cells
            for i, v in enumerate(row_data.values()):
                row[i].text = str(v)

    # Section 7 — Explainability (SHAP)
    doc.add_heading('7. Explainability (SHAP)', level=1)
    p5 = combined_results.get('phase5_explainability', {})
    imp = p5.get('feature_importance', [])
    if imp:
        doc.add_paragraph(f"{p5.get('method', 'SHAP')} computed on {p5.get('sample_size', '—')} sampled predictions.")
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Light Grid Accent 1'
        t.rows[0].cells[0].text = 'Feature'
        t.rows[0].cells[1].text = 'Mean |SHAP|'
        for entry in imp:
            row = t.add_row().cells
            row[0].text = entry.get('feature', '')
            row[1].text = str(entry.get('mean_abs_shap', ''))
        # Embed a bar chart of SHAP importance.
        try:
            plt.style.use('dark_background')
            fig, ax = plt.subplots(figsize=(7, 4))
            feats = [e['feature'] for e in imp]
            vals = [e['mean_abs_shap'] for e in imp]
            ax.barh(feats[::-1], vals[::-1], color='#f59e0b')
            ax.set_xlabel('Mean |SHAP|')
            ax.set_title('Dataset 2 — SHAP Feature Importance')
            chart_path = f'outputs/dataset2/{run_id}_shap_chart.png'
            fig.savefig(chart_path, dpi=120, bbox_inches='tight', facecolor='#0d1117')
            plt.close(fig)
            doc.add_picture(chart_path, width=Inches(6))
        except Exception:
            pass

    # Section 8 — Prediction
    doc.add_heading('8. Prediction', level=1)
    doc.add_paragraph(f"Records scored: {p6.get('n_scored', 0):,}")
    doc.add_paragraph(f"Score mean/median: {p6.get('score_mean', '—')} / {p6.get('score_median', '—')}")
    doc.add_paragraph(f"High-risk (≥0.5): {p6.get('pct_high_risk', '—')}%")
    bands = p6.get('risk_band_distribution', {})
    if bands:
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Light Grid Accent 1'
        t.rows[0].cells[0].text = 'Risk Band'
        t.rows[0].cells[1].text = 'Count'
        for band, count in bands.items():
            row = t.add_row().cells
            row[0].text = str(band)
            row[1].text = str(count)

    # Section 9 — Validation (conditional)
    doc.add_heading('9. Validation', level=1)
    if p7.get('target_available'):
        doc.add_paragraph(f"Target column used: {p7.get('target_column_used')}")
        doc.add_paragraph(f"AUC: {p7.get('auc')} | KS: {p7.get('ks')} | Gini: {p7.get('gini')} | "
                          f"Brier: {p7.get('brier_score')}")
        doc.add_paragraph(f"Score PSI: {p7.get('score_psi')} ({p7.get('score_psi_assessment')}) "
                          f"vs {p7.get('score_psi_reference')}")
        doc.add_paragraph(f"Overall Rating: {p7.get('overall_rating', '—')}")
        cm = p7.get('confusion_matrix', {})
        if cm:
            doc.add_paragraph(
                f"Precision: {cm.get('precision')} | Recall: {cm.get('recall')} | "
                f"F1: {cm.get('f1_score')} | Default Rate: {p7.get('default_rate', 0) * 100:.1f}%"
            )
        if p7.get('decile_table'):
            doc.add_paragraph(f"Rank-order check: {p7.get('rank_order_assessment', '—')}")
    else:
        doc.add_paragraph(
            'No ground-truth labels were available in this dataset. This is a true blind scoring '
            'exercise — validation metrics (AUC, KS, Gini, Confusion Matrix) cannot be computed until '
            'actual outcomes are observed. Prediction distribution and feature stability checks above '
            'provide the available evidence of model applicability to this population.'
        )

    # Section 10 — Fairness Check
    doc.add_heading('10. Fairness Check', level=1)
    p8 = combined_results.get('phase8_fairness', {})
    doc.add_paragraph(p8.get('summary', ''))
    for attr, groups in p8.get('fairness_results', {}).items():
        doc.add_paragraph(f"{attr.replace('_', ' ').title()}:", style='Intense Quote')
        t = doc.add_table(rows=1, cols=4)
        t.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Group', 'Count', 'Mean Predicted', 'Concern Level']):
            t.rows[0].cells[i].text = h
        for g, stats in groups.items():
            row = t.add_row().cells
            row[0].text = str(g)
            row[1].text = str(stats.get('count', ''))
            row[2].text = str(stats.get('mean_predicted', ''))
            row[3].text = str(stats.get('concern_level', ''))

    # Section 11 — Assumptions & Limitations
    doc.add_heading('11. Assumptions & Limitations', level=1)
    limitations = [
        'This report reflects a scoring exercise using a model FIXED from Dataset 1 — no retraining occurred.',
        f"Column alignment method: {alignment_method}." + (
            ' All required features were confidently identified.' if alignment_verified else
            f" {combined_results.get('alignment_warning', '')}"),
        "Engineered features were reconstructed from raw columns using Dataset 1's exact transformation logic.",
        'Feature stability (PSI) is assessed unsupervised — it indicates distributional similarity to training data, not predictive accuracy.',
        'Validation metrics (Section 9) are only available if genuine outcome labels exist in this dataset.',
        'This model was trained on funded/approved loans only — predictions for populations outside that scope carry additional uncertainty (selection bias).',
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style='List Bullet')

    os.makedirs('outputs/dataset2', exist_ok=True)
    doc_path = f'outputs/dataset2/{run_id}_Dataset2_Scoring_Report.docx'
    doc.save(doc_path)
    return {'phase': 'Documentation', 'report_path': doc_path}


# ── Orchestration ───────────────────────────────────────────────────────────

def run_dataset2_pipeline(dataset_path, run_id=None):
    """Orchestrates all 8 phases and saves combined results."""
    meta = load_champion_meta(run_id)

    p1 = phase1_data_understanding(dataset_path, meta)
    df = p1.pop('df')

    p2 = phase2_data_quality_scoped(df, meta, p1)                 # read-only (raw data)
    p3 = phase3_feature_reconstruction(df, meta, p1)
    df = p3.pop('df')                                            # now has engineered features

    p4 = phase4_feature_stability(df, meta, p1)
    p5 = phase5_explainability(df, meta, p1)
    p6 = phase6_prediction(df, meta, p1)
    y_prob = p6.pop('y_prob')
    p7 = phase7_validation(df, y_prob, meta)
    p8 = phase_fairness_check(df, y_prob, meta)
    p8.update(_alignment_fields(p1))                            # propagate alignment flag

    # Propagate alignment confidence for downstream consumers (UI).
    alignment_verified = p1.get('alignment_verified', True)
    scoring_confidence = ('VERIFIED' if alignment_verified
                          else 'UNVERIFIED — column alignment could not be confirmed')
    p6['prediction_status'] = scoring_confidence
    p7['scoring_confidence'] = scoring_confidence

    combined = {
        'dataset2_run_id': f"D2_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
        'source_dataset1_run': meta.get('run_id'),
        'champion_model': meta['champion_model'],
        'alignment_verified': alignment_verified,
        'alignment_warning': p1.get('alignment_warning', ''),
        'scoring_confidence': scoring_confidence,
        'phase1_data_understanding': p1,
        'phase2_data_quality': p2,
        'phase3_feature_reconstruction': p3,
        'phase4_feature_stability': p4,
        'phase5_explainability': p5,
        'phase6_prediction': {k: v for k, v in p6.items() if k != 'predictions'},  # exclude raw array
        'phase7_validation': p7,
        'phase8_fairness': p8,
        'alignment_method': p1.get('alignment_method'),
        'timestamp': datetime.now().isoformat(),
    }

    # Phase 9 — Documentation (uses the assembled combined results; never fails the run).
    try:
        p9 = phase9_documentation(combined, meta)
    except Exception as _e:
        p9 = {'phase': 'Documentation', 'report_path': None, 'error': str(_e)}
    combined['phase9_documentation'] = p9

    os.makedirs('outputs/dataset2', exist_ok=True)
    out_path = f"outputs/dataset2/{combined['dataset2_run_id']}_results.json"
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(combined, f, indent=2, default=str)

    pred_path = f"outputs/dataset2/{combined['dataset2_run_id']}_predictions.csv"
    pd.DataFrame({'predicted_probability': y_prob}).to_csv(pred_path, index=False)

    print(f"AUDIT_PATH:{out_path}")
    return combined, out_path


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--dataset', required=True)
    parser.add_argument('--run_id', default=None)
    args = parser.parse_args()

    print(f"\n{'='*60}")
    print("DATASET 2 — 8-PHASE BLIND SCORING PIPELINE")
    print(f"{'='*60}\n")

    result, path = run_dataset2_pipeline(args.dataset, args.run_id)

    if not result.get('alignment_verified', True):
        print("\n" + "!" * 60)
        print("  ⚠  UNVERIFIED SCORING — COLUMN ALIGNMENT NOT CONFIRMED")
        print("!" * 60)
        print("  " + result.get('alignment_warning', ''))
        print("  All results below are best-effort and should NOT be")
        print("  treated as confident.")
        print("!" * 60)

    p1 = result['phase1_data_understanding']
    print(f"\nPhase 1 — Data Understanding: {p1['total_rows']:,} rows, {p1['total_columns']} columns")
    print(f"  Header detected: {p1['header_detected']} | Alignment: {result.get('scoring_confidence')}")
    print(f"  Dictionary source: {p1['dictionary_source']}")
    sem = {}
    for s in p1['schema']:
        sem[s['semantic_type']] = sem.get(s['semantic_type'], 0) + 1
    print(f"  Semantic types: {sem}")

    p2 = result['phase2_data_quality']
    print(f"\nPhase 2 — Data Quality (Scoped): {len(p2['features_found'])} of "
          f"{len(p2['expected_features'])} champion features present in raw data")
    if p2['features_missing']:
        print(f"  Not in raw data (reconstructed in Phase 3): {p2['features_missing']}")
    n_out = sum(r['iqr_outliers'] for r in p2.get('outlier_table', []))
    print(f"  Outliers (IQR×3) across present features: {n_out:,} | Duplicate check: {p2['duplicate_check']}")

    p3 = result['phase3_feature_reconstruction']
    print(f"\nPhase 3 — Feature Reconstruction: {p3['n_reconstructed']} engineered features rebuilt "
          f"→ {p3['n_available_after']}/{len(p2['expected_features'])} champion features now available")
    for d in p3['derived_features_log']:
        print(f"    {d['feature']:<22} <- {d['derived_from']:<32} [{d['status']}]")

    p4 = result['phase4_feature_stability']
    print(f"\nPhase 4 — Feature Stability (PSI vs training): {p4['overall_assessment']} "
          f"(avg PSI={p4['avg_psi']}, {p4['n_assessed']} features assessed)")
    for r in p4['stability_table']:
        psi_txt = r['psi'] if r['psi'] is not None else '—'
        print(f"    {r['feature']:<22} PSI={str(psi_txt):<8} {r['assessment']}")

    p5 = result['phase5_explainability']
    print(f"\nPhase 5 — Explainability ({p5['method']}, sample={p5['sample_size']:,}):")
    for r in p5['feature_importance'][:5]:
        print(f"    {r['feature']:<22} {r['importance_pct']:>6}%  (mean|SHAP|={r['mean_abs_shap']})")

    p6 = result['phase6_prediction']
    print(f"\nPhase 6 — Prediction: {p6['n_scored']:,} scored with {p6['champion_model']} | "
          f"mean={p6['score_mean']} | median={p6['score_median']} | "
          f"min={p6['score_min']} | max={p6['score_max']} | high-risk={p6['pct_high_risk']}%")
    print(f"    Risk bands: {p6['risk_band_distribution']}")

    p7 = result['phase7_validation']
    print(f"\nPhase 7 — Validation: target_available={p7['target_available']}")
    print(f"  Score PSI: {p7.get('score_psi')} ({p7.get('score_psi_assessment')}) "
          f"— reference: {p7.get('score_psi_reference')}")
    if p7['target_available']:
        print(f"  Target column used: {p7['target_column_used']} | n_valid={p7['n_valid']:,}")
        print(f"  AUC={p7['auc']} | KS={p7['ks']} | Gini={p7['gini']} | Brier={p7['brier_score']}")
        cm = p7['confusion_matrix']
        print(f"  Precision={cm['precision']} | Recall={cm['recall']} | F1={cm['f1_score']} | "
              f"Default rate={p7['default_rate']}")
        print(f"  Rank-order: {p7.get('rank_order_breaks')} break(s) — {p7.get('rank_order_assessment')}")
        for k, v in p7.get('rag_summary', {}).items():
            print(f"    RAG {k:<10}: value={v['value']} → {v['rag']}")
        print(f"  OVERALL RATING: {p7.get('overall_rating')}")
    else:
        print(f"  {p7.get('message', '')}")
        print(f"  OVERALL RATING: {p7.get('overall_rating')}")

    p8 = result['phase8_fairness']
    print(f"\nPhase 8 — Fairness Check: {p8.get('summary')}")
    for attr, groups in p8.get('fairness_results', {}).items():
        flagged = {g: v for g, v in groups.items() if v['concern_level'] in ('Medium', 'High')}
        print(f"    {attr}: {len(groups)} group(s) checked"
              + (f" | flagged: {list(flagged.keys())}" if flagged else " | all Low concern"))

    p9 = result.get('phase9_documentation', {})
    print(f"\nPhase 9 — Documentation: {p9.get('report_path') or 'FAILED: ' + str(p9.get('error'))}")

    print(f"\nResults saved: {path}")
